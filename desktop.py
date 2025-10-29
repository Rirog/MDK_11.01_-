import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Inches
from datetime import datetime
import requests
import os

API_BASE_URL = "http://localhost:8000"


class CarTradingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AvtoLimonchik")
        self.root.geometry("1920x1080")
        self.root.configure(bg='#355c7d')
        self.root.resizable(False, False)

        self.center_window()
        
        self.auth_token = None
        self.current_user = None
        self.selected_user = None
        self.selected_stamp = None 
        self.selected_model = None 

        self.setup_styles()
        self.show_auth_frame()

    def center_window(self):
        """Центрирование окна на экране"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry('{}x{}+{}+{}'.format(width, height, x, y))
    
    def setup_styles(self):
        """Настройка стилей приложения"""
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        self.colors = {
            'primary': '#6c5b7b',
            'secondary': '#355c7d',
            'success': '#355c7d',
            'danger': '#e86375',
            'warning': '#ffd022',
            'light': "#f5e1d9",
            'dark': '#355c7d',
            'background': '#355c7d',
            'accent': '#355c7d'
        }
        
        self.style.configure('TFrame', background=self.colors['background'])
        self.style.configure('TLabel', background=self.colors['background'], 
                           foreground=self.colors['light'], font=('Arial', 10))
        self.style.configure('TButton', font=('Arial', 10, 'bold'))
        self.style.configure('Header.TLabel', font=('Arial', 20, 'bold'), 
                           foreground=self.colors['light'])
        self.style.configure('Accent.TButton', background=self.colors['accent'],
                           foreground='white', focuscolor='none')
        self.style.configure('Secondary.TButton', background=self.colors['dark'],
                           foreground=self.colors['light'], focuscolor='none')
        self.style.configure('Success.TButton', background=self.colors['success'],
                           foreground='white', focuscolor='none')
        self.style.configure('Danger.TButton', background=self.colors['danger'],
                           foreground='white', focuscolor='none')
        
        self.style.map('Accent.TButton', 
                      background=[('active', '#6c5b7b'),
                                 ('pressed', '#355c7d')])
        self.style.map('Secondary.TButton',
                      background=[('active', '#6c5b7b'),
                                 ('pressed', '#355c7d')])
        self.style.map('Success.TButton',
                      background=[('active', '#6c5b7b'),
                                 ('pressed', '#355c7d')])
        self.style.map('Danger.TButton',
                      background=[('active', '#6c5b7b'),
                                 ('pressed', '#355c7d')])
    
    def clear_window(self):
        """Очистка окна"""
        for widget in self.root.winfo_children():
            widget.destroy()
    
    def create_card_frame(self, parent):
        """Создать карточку для контента"""
        card = tk.Frame(parent, bg=self.colors['light'], relief='raised', 
                       bd=1, padx=20, pady=20)
        return card

    def show_filter_sort_options(self, parent, filter_callback):
        """фильтрация и сортировка с кнопкой экспорта"""
        filter_frame = tk.Frame(parent, bg=self.colors['light'])
        filter_frame.pack(fill=tk.X, pady=(0, 15))

        main_container = tk.Frame(filter_frame, bg=self.colors['light'])
        main_container.pack(fill=tk.X)

        filter_label = tk.Label(main_container, text="🔍 Фильтры и сортировка", 
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 11, 'bold'), anchor='w')
        filter_label.pack(fill=tk.X, pady=(0, 8))

        controls_frame = tk.Frame(main_container, bg=self.colors['light'])
        controls_frame.pack(fill=tk.X)

        stamp_frame = tk.Frame(controls_frame, bg=self.colors['light'])
        stamp_frame.pack(side=tk.LEFT, padx=(0, 10))
        
        stamp_label = tk.Label(stamp_frame, text="Марка:", bg=self.colors['light'], 
                            fg=self.colors['dark'], font=('Arial', 9))
        stamp_label.pack(anchor='w')
        
        self.filter_stamp_combo = ttk.Combobox(stamp_frame, font=('Arial', 9), width=12)
        self.filter_stamp_combo.pack(pady=(2, 0))
        self.filter_stamp_combo.bind('<<ComboboxSelected>>', self.on_stamp_selected)

        model_frame = tk.Frame(controls_frame, bg=self.colors['light'])
        model_frame.pack(side=tk.LEFT, padx=(0, 10))
        
        model_label = tk.Label(model_frame, text="Модель:", bg=self.colors['light'], 
                            fg=self.colors['dark'], font=('Arial', 9))
        model_label.pack(anchor='w')
        
        self.filter_model_combo = ttk.Combobox(model_frame, font=('Arial', 9), width=12)
        self.filter_model_combo.pack(pady=(2, 0))

        price_frame = tk.Frame(controls_frame, bg=self.colors['light'])
        price_frame.pack(side=tk.LEFT, padx=(0, 10))
        
        price_label = tk.Label(price_frame, text="Цена до:", bg=self.colors['light'], 
                            fg=self.colors['dark'], font=('Arial', 9))
        price_label.pack(anchor='w')
        
        self.filter_price_entry = ttk.Entry(price_frame, font=('Arial', 9), width=10)
        self.filter_price_entry.pack(pady=(2, 0))

        sort_frame = tk.Frame(controls_frame, bg=self.colors['light'])
        sort_frame.pack(side=tk.LEFT, padx=(0, 10))
        
        sort_label = tk.Label(sort_frame, text="Сортировка:", bg=self.colors['light'], 
                            fg=self.colors['dark'], font=('Arial', 9))
        sort_label.pack(anchor='w')
        
        self.sort_combo = ttk.Combobox(sort_frame, font=('Arial', 9), width=15,
                                    values=["Цена (по возрастанию)", "Цена (по убыванию)", 
                                            "Пробег (по возрастанию)", "Пробег (по убыванию)"])
        self.sort_combo.set("Цена (по возрастанию)")
        self.sort_combo.pack(pady=(2, 0))

        buttons_frame = tk.Frame(controls_frame, bg=self.colors['light'])
        buttons_frame.pack(side=tk.LEFT, padx=(10, 0))
        
        apply_btn = ttk.Button(buttons_frame, text="Применить", 
                            style='Accent.TButton',
                            command=filter_callback,
                            width=12)
        apply_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        reset_btn = ttk.Button(buttons_frame, text="Сбросить", 
                            style='Secondary.TButton',
                            command=self.reset_filters,
                            width=12)
        reset_btn.pack(side=tk.LEFT, padx=(0, 5))

        export_btn = ttk.Button(buttons_frame, text="📄 Экспорт в DOCX",
                            style='Success.TButton',
                            command=self.export_filtered_cars_to_docx,
                            width=30)
        export_btn.pack(side=tk.LEFT)
        
        return filter_frame

    def export_filtered_cars_to_docx(self):
        """Экспорт отфильтрованных автомобилей в DOCX"""
        try:
            if not hasattr(self, 'filtered_cars') or not self.filtered_cars:
                messagebox.showinfo("Информация", "Нет данных для экспорта")
                return

            doc = Document()

            title = doc.add_heading('Отфильтрованный список автомобилей AvtoLimonchik', 0)

            current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            doc.add_paragraph(f'Дата формирования отчета: {current_time}')

            filters_info = doc.add_paragraph()
            filters_info.add_run('Примененные фильтры:\n').bold = True
            
            filter_text = ""
            if hasattr(self, 'filter_stamp_combo') and self.filter_stamp_combo.get():
                filter_text += f"Марка: {self.filter_stamp_combo.get()}\n"
            if hasattr(self, 'filter_model_combo') and self.filter_model_combo.get():
                filter_text += f"Модель: {self.filter_model_combo.get()}\n"
            if hasattr(self, 'filter_price_entry') and self.filter_price_entry.get():
                filter_text += f"Цена до: {self.filter_price_entry.get()} руб\n"
            if hasattr(self, 'sort_combo') and self.sort_combo.get():
                filter_text += f"Сортировка: {self.sort_combo.get()}\n"
            
            if filter_text:
                filters_info.add_run(filter_text)
            else:
                filters_info.add_run('Фильтры не применены\n')

            stats = doc.add_paragraph()
            stats.add_run('Статистика:\n').bold = True
            stats.add_run(f'Найдено автомобилей: {len(self.filtered_cars)} шт.\n')
            
            if self.filtered_cars:
                total_price = sum(car.get('price', 0) for car in self.filtered_cars)
                avg_price = total_price / len(self.filtered_cars)
                stats.add_run(f'Средняя цена: {avg_price:,.0f} руб.\n'.replace(",", " "))
                stats.add_run(f'Общая стоимость: {total_price:,} руб.\n'.replace(",", " "))

            doc.add_heading('Список автомобилей', level=1)
            
            if self.filtered_cars:
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '№'
                hdr_cells[1].text = 'Марка'
                hdr_cells[2].text = 'Модель'
                hdr_cells[3].text = 'Пробег (км)'
                hdr_cells[4].text = 'Цена (руб)'
                hdr_cells[5].text = 'VIN'

                for i, car in enumerate(self.filtered_cars, 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = car.get('stamp', '')
                    row_cells[2].text = car.get('model', '')
                    row_cells[3].text = f"{car.get('run_km', 0):,}".replace(",", " ")
                    row_cells[4].text = f"{car.get('price', 0):,}".replace(",", " ")
                    row_cells[5].text = car.get('vin', '')
            else:
                doc.add_paragraph('Нет автомобилей, соответствующих фильтрам')

            reports_dir = "reports"
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)

            filename = f"filtered_cars_AvtoLimonchik_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(reports_dir, filename)
            doc.save(filepath)
            
            messagebox.showinfo("Успех", f"Отфильтрованный список автомобилей успешно экспортирован в файл:\n{filepath}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте: {str(e)}")

    def apply_car_filters(self, cars):
        """Применить фильтры и сортировку к списку автомобилей"""
        filtered_cars = cars.copy()

        if hasattr(self, 'filter_stamp_combo') and self.filter_stamp_combo.get():
            selected_stamp = self.filter_stamp_combo.get()
            filtered_cars = [car for car in filtered_cars 
                            if car.get('stamp') == selected_stamp]

        if hasattr(self, 'filter_model_combo') and self.filter_model_combo.get():
            selected_model = self.filter_model_combo.get()
            filtered_cars = [car for car in filtered_cars 
                            if car.get('model') == selected_model]

        if hasattr(self, 'filter_price_entry') and self.filter_price_entry.get():
            try:
                max_price = int(self.filter_price_entry.get())
                filtered_cars = [car for car in filtered_cars 
                            if car.get('price', 0) <= max_price]
            except ValueError:
                pass

        if hasattr(self, 'sort_combo'):
            sort_option = self.sort_combo.get()
            if sort_option == "Цена (по возрастанию)":
                filtered_cars.sort(key=lambda x: x.get('price', 0))
            elif sort_option == "Цена (по убыванию)":
                filtered_cars.sort(key=lambda x: x.get('price', 0), reverse=True)
            elif sort_option == "Пробег (по возрастанию)":
                filtered_cars.sort(key=lambda x: x.get('run_km', 0))
            elif sort_option == "Пробег (по убыванию)":
                filtered_cars.sort(key=lambda x: x.get('run_km', 0), reverse=True)

        self.filtered_cars = filtered_cars
        
        return filtered_cars

    def on_stamp_selected(self, event=None):
        """Обработчик выбора марки - обновляет список моделей"""
        selected_stamp = self.filter_stamp_combo.get()
        if selected_stamp and hasattr(self, 'original_cars_data'):
            models = list(set(car.get('model', '') for car in self.original_cars_data 
                            if car.get('stamp') == selected_stamp and car.get('model')))
            self.filter_model_combo['values'] = models
            self.filter_model_combo.set('')

    def reset_filters(self):
        """Сбросить фильтры"""
        if hasattr(self, 'filter_stamp_combo'):
            self.filter_stamp_combo.set('')
        if hasattr(self, 'filter_model_combo'):
            self.filter_model_combo.set('')
        if hasattr(self, 'filter_price_entry'):
            self.filter_price_entry.delete(0, tk.END)
        if hasattr(self, 'sort_combo'):
            self.sort_combo.set("Цена (по возрастанию)")

        if hasattr(self, 'current_filter_callback'):
            self.current_filter_callback()

    def export_my_purchases_report(self):
        """Экспорт истории покупок пользователя в DOCX"""
        try:
            purchases = self.get_my_purchases_data()
            
            if not purchases:
                messagebox.showinfo("Информация", "У вас нет покупок для экспорта")
                return

            doc = Document()

            title = doc.add_heading('Мои покупки в AvtoLimonchik', 0)

            current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            doc.add_paragraph(f'Дата формирования отчета: {current_time}')
            doc.add_paragraph(f'Пользователь: {self.current_user.get("name", "")}')
            doc.add_paragraph(f'Email: {self.current_user.get("email", "")}')
            doc.add_paragraph(f'Телефон: {self.current_user.get("phone", "")}')

            doc.add_heading('Общая статистика', level=1)
            total_cars = len(purchases)
            total_spent = sum(purchase.get('price', 0) for purchase in purchases)
            
            stats = doc.add_paragraph()
            stats.add_run('Общее количество автомобилей: ').bold = True
            stats.add_run(f'{total_cars} шт.\n')
            stats.add_run('Общая сумма покупок: ').bold = True
            stats.add_run(f'{total_spent:,} руб.\n'.replace(",", " "))
            
            if total_cars > 0:
                avg_price = total_spent / total_cars
                stats.add_run('Средняя стоимость автомобиля: ').bold = True
                stats.add_run(f'{avg_price:,.0f} руб.\n'.replace(",", " "))

            doc.add_heading('Детальная информация о покупках', level=1)
            
            for i, purchase in enumerate(purchases, 1):
                car_info = purchase.get('car', {})
                purchase_date = purchase.get('date_buy', '')[:10]

                purchase_heading = doc.add_heading(f'Покупка #{i}', level=2)

                car_table = doc.add_table(rows=6, cols=2)
                car_table.style = 'Table Grid'
 
                cells_data = [
                    ('Дата покупки:', purchase_date),
                    ('Марка:', car_info.get('stamp', 'Не указана')),
                    ('Модель:', car_info.get('model', 'Не указана')),
                    ('VIN:', car_info.get('vin', 'Не указан')),
                    ('Пробег:', f"{car_info.get('run_km', 0):,} км".replace(",", " ")),
                    ('Стоимость:', f"{purchase.get('price', 0):,} руб".replace(",", " "))
                ]
                
                for row, (label, value) in enumerate(cells_data):
                    car_table.cell(row, 0).text = label
                    car_table.cell(row, 1).text = value

                description = car_info.get('description')
                if description:
                    doc.add_paragraph('Описание автомобиля:')
                    desc_para = doc.add_paragraph(description)
                    desc_para.style = 'List Bullet'

                if i < len(purchases):
                    doc.add_paragraph()
                    doc.add_paragraph('―' * 50)
                    doc.add_paragraph()

            reports_dir = "reports"
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)

            filename = f"my_purchases_AvtoLimonchik{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(reports_dir, filename)
            doc.save(filepath)
            
            messagebox.showinfo("Успех", f"Отчет о покупках успешно экспортирован в файл:\n{filepath}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте отчета: {str(e)}")

    def get_my_purchases_data(self):
        """Получить данные о покупках пользователя"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/my_purchases", headers=headers)
            
            if response.status_code == 200:
                return response.json()
            elif response.status_code == 404:
                return []
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки покупок")
                print(f"Ошибка: {error_msg}")
                return []
                
        except requests.exceptions.RequestException as e:
            print(f"Ошибка подключения: {e}")
            return []
        
    def collect_firm_statistics_for_export(self):
        """Сбор статистики фирмы для экспорта"""
        stats = {
            'available_cars': 0,
            'firm_purchased_cars': 0, 
            'firm_solds_cars': 0,  
            'total_revenue': 0,
            'total_purchase_cost': 0, 
            'recent_firm_sales': [],
            'recent_firm_purchases': [],
            'report_date': datetime.now().strftime("%d.%m.%Y %H:%M"),
            'firm_name': 'AvtoLimonchik'
        }
        
        try:
            headers = {"token": self.auth_token}

            response = requests.get(f"{API_BASE_URL}/users/cars/available", headers=headers)
            if response.status_code == 200:
                available_cars = response.json()
                stats['available_cars'] = len(available_cars)

            response = requests.get(f"{API_BASE_URL}/admin/sales/", headers=headers)
            if response.status_code == 200:
                firm_sales = response.json()
                stats['firm_purchased_cars'] = len(firm_sales)

                total_revenue = 0
                for sale in firm_sales:
                    price = sale.get('price', 0)
                    total_revenue += price
                
                stats['total_revenue'] = total_revenue

                recent_sales = sorted(firm_sales, key=lambda x: x.get('date_sale', ''), reverse=True)[:10]
                for sale in recent_sales:
                    stats['recent_firm_sales'].append({
                        'buyer_name': sale.get('buyer', {}).get('name', 'Неизвестно'),
                        'car_info': f"{sale.get('car', {}).get('stamp', '')} {sale.get('car', {}).get('model', '')}",
                        'price': sale.get('price', 0),
                        'date': sale.get('date_sale', '')[:10]
                    })
            response = requests.get(f"{API_BASE_URL}/admin/shopping/", headers=headers)
            if response.status_code == 200:
                firm_purchases = response.json()
                stats['firm_solds_cars'] = len(firm_purchases)

                total_purchase_cost = 0
                for purchase in firm_purchases:
                    price = purchase.get('price', 0)
                    total_purchase_cost += price
                
                stats['total_purchase_cost'] = total_purchase_cost

                recent_purchases = sorted(firm_purchases, key=lambda x: x.get('date_buy', ''), reverse=True)[:10]
                for purchase in recent_purchases:
                    stats['recent_firm_purchases'].append({
                        'buyer_name': purchase.get('buyer', {}).get('name', 'Неизвестно'),
                        'car_info': f"{purchase.get('car', {}).get('stamp', '')} {purchase.get('car', {}).get('model', '')}",
                        'price': purchase.get('price', 0),
                        'date': purchase.get('date_buy', '')[:10]
                    })
            
        except Exception as e:
            print(f"Ошибка при сборе статистики фирмы: {e}")
        
        return stats

    def export_firm_report(self):
        """Экспорт отчета по деятельности фирмы в DOCX"""
        try:

            doc = Document()

            title = doc.add_heading('Отчет о деятельности AvtoLimonchik', 0)
            current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            doc.add_paragraph(f'Дата формирования: {current_time}')
            doc.add_paragraph(f'Отчет подготовил: {self.current_user.get("name", "")}')

            stats = self.collect_firm_statistics_for_export()

            doc.add_heading('Общая статистика деятельности фирмы', level=1)
            
            general_stats = doc.add_paragraph()
            general_stats.add_run('Автомобили в продаже: ').bold = True
            general_stats.add_run(f'{stats["available_cars"]} ед.\n')
            
            general_stats.add_run('Машин продано фирмой: ').bold = True
            general_stats.add_run(f'{stats["firm_solds_cars"]} ед.\n')
            
            general_stats.add_run('Машин куплено фирмой: ').bold = True
            general_stats.add_run(f'{stats["firm_purchased_cars"]} ед.\n')
            
            general_stats.add_run('Общая выручка от продаж: ').bold = True
            general_stats.add_run(f'{stats["total_purchase_cost"]:,} руб.\n'.replace(",", " "))
            
            general_stats.add_run('Общие затраты на покупки: ').bold = True
            general_stats.add_run(f'{stats["total_revenue"]:,} руб.\n'.replace(",", " "))
        
            profit = stats["total_purchase_cost"] - stats["total_revenue"]
            profit_status = "Прибыль" if profit >= 0 else "Убыток"
            
            general_stats.add_run(f'Финансовый результат: ').bold = True
            general_stats.add_run(f'{profit_status}: {abs(profit):,} руб.'.replace(",", " "))

            doc.add_heading('Последние покупки фирмы', level=1)
            if stats['recent_firm_sales']:
                for sale in stats['recent_firm_sales']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{sale["date"]} - {sale["buyer_name"]} - {sale["car_info"]} - {sale["price"]:,} руб'.replace(",", " "))
            else:
                doc.add_paragraph('Нет данных о покупках')

            doc.add_heading('Последние продажи фирмы', level=1)
            if stats['recent_firm_purchases']:
                for purchase in stats['recent_firm_purchases']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{purchase["date"]} - {purchase["buyer_name"]} - {purchase["car_info"]} - {purchase["price"]:,} руб'.replace(",", " "))
            else:
                doc.add_paragraph('Нет данных о продажах')

            doc.add_heading('Финансовая сводка', level=1)
            financial_info = doc.add_paragraph()
            financial_info.add_run('Выручка от продаж: ').bold = True
            financial_info.add_run(f'{stats["total_purchase_cost"]:,} руб.\n'.replace(",", " "))
            
            financial_info.add_run('Затраты на закупки: ').bold = True
            financial_info.add_run(f'{stats["total_revenue"]:,} руб.\n'.replace(",", " "))
            
            financial_info.add_run('Чистый финансовый результат: ').bold = True
            financial_info.add_run(f'{profit_status}: {abs(profit):,} руб.\n'.replace(",", " "))


            reports_dir = "reports"
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)
                
            filename = f"firm_report_AvtoLimonchik{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(reports_dir, filename)
            doc.save(filepath)
            
            messagebox.showinfo("Успех", f"Отчет о деятельности фирмы успешно экспортирован в файл:\n{filepath}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте отчета: {str(e)}")

    def show_auth_frame(self):
        """Показать фрейм авторизации"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(10, 20))

        car_icon = tk.Label(header_frame, text="🚗", font=('Arial', 40), 
                           bg=self.colors['background'], fg=self.colors['light'])
        car_icon.pack(pady=(0, 10))

        header_label = ttk.Label(header_frame, text="AvtoLimonchik", 
                               style='Header.TLabel')
        header_label.pack(pady=(0, 5))

        sub_label = ttk.Label(header_frame, text="Платформа для покупки и продажи автомобилей", 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        separator = ttk.Separator(header_frame, orient='horizontal')
        separator.pack(fill=tk.X, pady=10)

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)

        welcome_text = (
            "Добро пожаловать в AvtoLimonchik!\n\n"
            "Здесь вы можете:\n"
            "• Продать свой автомобиль\n" 
            "• Купить автомобиль у фирмы\n"
            "• Просматривать доступные автомобили\n\n"
            "Выберите действие для продолжения:"
        )
        
        welcome_label = tk.Label(card, text=welcome_text, 
                               bg=self.colors['light'], fg=self.colors['dark'],
                               font=('Arial', 11), justify=tk.LEFT)
        welcome_label.pack(pady=20, anchor='w')

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=20)

        login_btn = ttk.Button(btn_frame, text="Вход в систему", 
                              style='Accent.TButton', 
                              command=self.show_login)
        login_btn.pack(fill=tk.X, pady=8, ipady=10)

        register_btn = ttk.Button(btn_frame, text="Создать аккаунт", 
                                 style='Success.TButton', 
                                 command=self.show_register)
        register_btn.pack(fill=tk.X, pady=8, ipady=10)

        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(fill=tk.X, pady=(18, 10))
        
        footer_label = ttk.Label(footer_frame, text="© 2025 AvtoLimonchik. Все права защищены.", 
                               font=('Arial', 9), foreground='#95a5a6')
        footer_label.pack()
    
    def show_login(self):
        """Показать форму входа"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Вход в систему", 
                               style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        sub_label = ttk.Label(header_frame, text="Введите ваши учетные данные", 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(expand=True, pady=10)

        fields = [
            ("Электронная почта или телефон", "login_email_phone"),
            ("Пароль", "login_password")
        ]
        
        self.login_entries = {}
        
        for i, (label, field_name) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                         fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                         anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_container, font=('Arial', 11))
            if "password" in field_name:
                entry.config(show="•")
            entry.pack(fill=tk.X, pady=2, ipady=6)
            self.login_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(10, 10))

        login_btn = ttk.Button(btn_frame, text="Войти", 
                              style='Accent.TButton', 
                              command=self.perform_login)
        login_btn.pack(fill=tk.X, pady=6, ipady=8)

        back_btn = ttk.Button(btn_frame, text="Назад", 
                             style='Secondary.TButton', 
                             command=self.show_auth_frame)
        back_btn.pack(fill=tk.X, pady=6, ipady=6)
    
    def show_register(self):
        """Показать форму регистрации"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Регистрация", 
                               style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        sub_label = ttk.Label(header_frame, text="Создайте новый аккаунт", 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)
        
        form_container = tk.Frame(card, bg=self.colors['light'])
        form_container.pack(expand=True, pady=10)

        fields = [
            ("Электронная почта", "register_email"),
            ("Телефон", "register_phone"), 
            ("Полное имя", "register_full_name"),
            ("Пароль", "register_password")
        ]

        self.register_entries = {}
        
        for i, (label, field_name) in enumerate(fields):

            field_frame = tk.Frame(form_container, bg=self.colors['light'])
            field_frame.pack(fill=tk.X, pady=10)

            lbl = tk.Label(field_frame, text=label, 
                         bg=self.colors['light'], fg=self.colors['dark'],
                         font=('Arial', 10, 'bold'), anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_frame, font=('Arial', 11))
            if "password" in field_name:
                entry.config(show="•")
            entry.pack(fill=tk.X, ipady=5)
            self.register_entries[field_name] = entry

        separator = ttk.Separator(form_container, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        btn_frame = tk.Frame(form_container, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=10)

        register_btn = ttk.Button(btn_frame, text="Зарегистрироваться", 
                                 style='Success.TButton', 
                                 command=self.perform_register)
        register_btn.pack(fill=tk.X, pady=4, ipady=5)

        back_btn = ttk.Button(btn_frame, text="Назад", 
                             style='Secondary.TButton', 
                             command=self.show_auth_frame)
        back_btn.pack(fill=tk.X, pady=1, ipady=1)

        hint_label = tk.Label(form_container, 
                             text="После регистрации вы сможете войти в систему",
                             bg=self.colors['light'], fg='#7f8c8d',
                             font=('Arial', 9))
        hint_label.pack(pady=10)
    
    def perform_login(self):
        """Выполнить вход"""
        email_phone = self.login_entries["login_email_phone"].get().strip()
        password = self.login_entries["login_password"].get()
        
        if not email_phone or not password:
            messagebox.showerror("Ошибка", "Заполните все поля")
            return

        login_data = {"password": password}
        if "@" in email_phone:
            login_data["email"] = email_phone
        else:
            login_data["phone"] = email_phone
        
        try:
            response = requests.post(f"{API_BASE_URL}/users/auth/", json=login_data)
            
            if response.status_code == 200:
                data = response.json()
                self.auth_token = data["token"]
                messagebox.showinfo("Успех", "Вход выполнен успешно!")
                self.get_user_profile()
            else:
                error_msg = response.json().get("detail", "Ошибка авторизации")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")
    
    def perform_register(self):
        """Выполнить регистрацию"""
        data = {
            "email": self.register_entries["register_email"].get().strip(),
            "phone": self.register_entries["register_phone"].get().strip(),
            "full_name": self.register_entries["register_full_name"].get().strip(),
            "password": self.register_entries["register_password"].get()
        }

        for field, value in data.items():
            if not value:
                messagebox.showerror("Ошибка", "Заполните все поля")
                return
        
        try:
            response = requests.post(f"{API_BASE_URL}/users/register/", json=data)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Регистрация выполнена успешно!")
                for entry in self.register_entries.values():
                    entry.delete(0, tk.END)
                self.show_login()
            else:
                error_msg = response.json().get("detail", "Ошибка регистрации")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def get_user_profile(self):
        """Получить профиль пользователя"""
        if not self.auth_token:
            messagebox.showerror("Ошибка", "Токен авторизации отсутствует")
            return
        
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/me/", headers=headers)
            
            if response.status_code == 200:
                self.current_user = response.json()
                try:
                    role_response = requests.get(
                        f"{API_BASE_URL}/admin/users/{self.current_user['id']}", 
                        headers=headers
                    )
                    if role_response.status_code == 200:
                        user_data = role_response.json()
                        self.current_user['role'] = user_data.get('role', 'Пользователь')
                    else:
                        self.current_user['role'] = 'Пользователь'
                except Exception as e:
                    print(f"Ошибка при получении роли: {e}")
                    self.current_user['role'] = 'Пользователь'
                
                self.show_main_menu()
            else:
                error_msg = response.json().get("detail", "Ошибка получения профиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")


    def show_main_menu(self):
        """Показать главное меню"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Главное меню", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        user_role = self.current_user.get('role', 'Пользователь')
        role_text = f"Добро пожаловать в AvtoLimonchik ({user_role})"
        sub_label = ttk.Label(header_frame, text=role_text, 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)

        welcome_text = f"👋 Приветствуем, {self.current_user.get('name', 'Пользователь')}!"
        welcome_label = tk.Label(card, text=welcome_text,
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 14, 'bold'), anchor='w')
        welcome_label.pack(fill=tk.X, pady=(0, 20))

        menu_frame = tk.Frame(card, bg=self.colors['light'])
        menu_frame.pack(expand=True, pady=10)

        if self.current_user.get('role') == 'Администратор':
            menu_buttons = [
                    ("👤 Управление профилем", self.show_my_profil, 'Accent.TButton'),
                    ("🚗 Управление автомобилями", self.show_admin_cars_management, 'Success.TButton'),
                    ("🏷️ Управление марками", self.show_stamp_management, 'Secondary.TButton'),
                    ("🚙 Управление моделями", self.show_model_management, 'Secondary.TButton'),
                    ("📋 Управление анкетами", self.show_admin_anketi_management, 'Secondary.TButton'),
                    ("👥 Управление пользователями", self.show_user_management, 'Danger.TButton'),
                    ("🛒 История покупок", self.show_admin_purchases, 'Secondary.TButton'),
                    ("💰 История продаж", self.show_admin_sales, 'Secondary.TButton'),
                    ("📊 Отчет в DOCX", self.export_firm_report, 'Secondary.TButton'),
            ]
        else:

            menu_buttons = [
                ("👤 Управление профилем", self.show_my_profil, 'Accent.TButton'),
                ("🚗 Доступные автомобили", self.show_available_cars, 'Success.TButton'),
                ("📝 Мои анкеты", self.show_my_anketi, 'Secondary.TButton'),
                ("🛒 Мои покупки", self.show_my_purchases, 'Secondary.TButton'),
            ]

        self.menu_buttons = {}
        for i, (text, command, style) in enumerate(menu_buttons):
            btn_frame = tk.Frame(menu_frame, bg=self.colors['light'])
            btn_frame.pack(fill=tk.X, pady=8)
            
            btn = ttk.Button(btn_frame, text=text, style=style, command=command, width=30)
            btn.pack(ipady=12, anchor='center')
            self.menu_buttons[text] = btn

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        logout_frame = tk.Frame(card, bg=self.colors['light'])
        logout_frame.pack(fill=tk.X, pady=10)
        
        logout_btn = ttk.Button(logout_frame, text="🚪 Выйти из системы", 
                            style='Danger.TButton', 
                            command=self.confirm_logout,
                            width=30)
        logout_btn.pack(ipady=12, anchor='center')

        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(fill=tk.X, pady=(18, 10))
        
        footer_label = ttk.Label(footer_frame, text="© 2025 AvtoLimonchik. Все права защищены.", 
                            font=('Arial', 9), foreground='#95a5a6')
        footer_label.pack()
    

    def show_my_profil(self):
        """Профиль"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Управление профилем", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)

        profile_header = tk.Label(card, text="👤 Мой профиль", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        profile_header.pack(fill=tk.X, pady=(0, 15))

        user_info_frame = tk.Frame(card, bg=self.colors['light'])
        user_info_frame.pack(fill=tk.X, pady=10)

        user_fields = [
            ("Имя:", self.current_user.get('name', '')),
            ("Email:", self.current_user.get('email', '')),
            ("Телефон:", self.current_user.get('phone', '')),
            ("Роль:", self.current_user.get('role', 'Пользователь'))  
        ]
        
        for label, value in user_fields:
            field_frame = tk.Frame(user_info_frame, bg=self.colors['light'])
            field_frame.pack(fill=tk.X, pady=8)
            
            lbl = tk.Label(field_frame, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        width=10, anchor='w')
            lbl.pack(side=tk.LEFT)
            
            value_lbl = tk.Label(field_frame, text=value, bg=self.colors['light'], 
                            fg='#2c3e50', font=('Arial', 10),
                            anchor='w')
            value_lbl.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        profile_actions_frame = tk.Frame(card, bg=self.colors['light'])
        profile_actions_frame.pack(fill=tk.X, pady=10)

        profile_buttons = [
            ("✏️ Изменить профиль", self.show_edit_profile, 'Accent.TButton'),
            ("🔒 Сменить пароль", self.show_change_password, 'Secondary.TButton'),
        ]

        if self.current_user.get('role') != 'Администратор':
            profile_buttons.append(("🗑️ Удалить аккаунт", self.confirm_delete_account, 'Danger.TButton'))
        
        for text, command, style in profile_buttons:
            btn_frame = tk.Frame(profile_actions_frame, bg=self.colors['light'])
            btn_frame.pack(fill=tk.X, pady=6)
            
            btn = ttk.Button(btn_frame, text=text, style=style, command=command, width=25)
            btn.pack(ipady=8, anchor='center')
    

    def show_edit_profile(self):
        """Редактирования профиля"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование профиля", 
                               style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        sub_label = ttk.Label(header_frame, text="Измените данные вашего профиля", 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(expand=True, pady=10)

        fields = [
            ("Полное имя", "edit_full_name", self.current_user.get('name', '')),
            ("Телефон", "edit_phone", self.current_user.get('phone', ''))
        ]
        
        self.edit_entries = {}
        
        for i, (label, field_name, current_value) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)
            
            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                         fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                         anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_container, font=('Arial', 11))
            entry.insert(0, current_value)
            entry.pack(fill=tk.X, pady=2, ipady=6)
            self.edit_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                             style='Success.TButton', 
                             command=self.perform_edit_profile)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                              style='Secondary.TButton', 
                              command=self.show_main_menu)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)


    def show_change_password(self):
        """Смена пароля"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Смена пароля", 
                               style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        sub_label = ttk.Label(header_frame, text="Введите новый пароль", 
                            font=('Arial', 12), foreground='#bdc3c7')
        sub_label.pack(pady=(0, 10))

        card = self.create_card_frame(main_frame)
        card.pack(expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(expand=True, pady=10)

        fields = [
            ("Новый пароль", "new_password"),
            ("Подтвердите пароль", "confirm_password")
        ]
        
        self.password_entries = {}
        
        for i, (label, field_name) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                         fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                         anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_container, font=('Arial', 11), show="•")
            entry.pack(fill=tk.X, pady=2, ipady=6)
            self.password_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="🔒 Сменить пароль", 
                             style='Success.TButton', 
                             command=self.perform_change_password)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                              style='Secondary.TButton', 
                              command=self.show_main_menu)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)


    def perform_edit_profile(self):
        """Обновление профиля"""
        full_name = self.edit_entries["edit_full_name"].get().strip()
        phone = self.edit_entries["edit_phone"].get().strip()
        
        if not full_name or not phone:
            messagebox.showerror("Ошибка", "Заполните все поля")
            return
        
        try:
            headers = {"token": self.auth_token}
            data = {
                "full_name": full_name,
                "phone": phone
            }
            
            response = requests.put(f"{API_BASE_URL}/users/me/", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Профиль успешно обновлен!")
                self.get_user_profile() 
            else:
                error_msg = response.json().get("detail", "Ошибка обновления профиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def perform_change_password(self):
        """Смена пароля"""
        new_password = self.password_entries["new_password"].get()
        confirm_password = self.password_entries["confirm_password"].get()
        
        if not new_password or not confirm_password:
            messagebox.showerror("Ошибка", "Заполните все поля")
            return
        
        if new_password != confirm_password:
            messagebox.showerror("Ошибка", "Пароли не совпадают")
            return
        
        try:
            headers = {"token": self.auth_token}
            data = {"password": new_password}
            
            response = requests.put(f"{API_BASE_URL}/users/me/password", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Пароль успешно изменен!")
                self.show_main_menu()
            else:
                error_msg = response.json().get("detail", "Ошибка смены пароля")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_delete_account(self):
        """Подтверждение удаления аккаунта"""
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            "Вы уверены, что хотите удалить свой аккаунт?\n\nЭто действие нельзя отменить!",
            icon='warning'
        )
        
        if result:
            self.perform_delete_account()

    def perform_delete_account(self):
        """Удаление аккаунта"""
        try:
            headers = {"token": self.auth_token}
            response = requests.delete(f"{API_BASE_URL}/users/delete_me/", headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Аккаунт успешно удален!")
                self.auth_token = None
                self.current_user = None
                self.show_auth_frame()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления аккаунта")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_logout(self):
        """Подтверждение выхода из системы"""
        result = messagebox.askyesno(
            "Подтверждение выхода", 
            "Вы уверены, что хотите выйти из системы?",
            icon='question'
        )
        
        if result:
            self.perform_logout()

    def perform_logout(self):
        """Выполнить выход из системы"""
        try:
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/users/logout/", headers=headers)

            self.auth_token = None
            self.current_user = None
            self.show_auth_frame()
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Вы успешно вышли из системы!")
            else:
                pass
                
        except requests.exceptions.RequestException:
            self.auth_token = None
            self.current_user = None
            self.show_auth_frame()


    def show_available_cars(self):
        """Показать раздел доступных автомобилей"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Доступные автомобили", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        cars_header = tk.Label(card, text="🚗 Автомобили в продаже", 
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 14, 'bold'), anchor='w')
        cars_header.pack(fill=tk.X, pady=(0, 15))
        self.current_filter_callback = lambda: self.refresh_cars_list(card)
        filter_frame = self.show_filter_sort_options(card, self.current_filter_callback)
        
        cars_list_frame = tk.Frame(card, bg=self.colors['light'])
        cars_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_available_cars(cars_list_frame)

    def refresh_cars_list(self, parent_card):
        """Обновить список автомобилей"""
        for widget in parent_card.winfo_children():
            if isinstance(widget, tk.Frame) and widget != parent_card.winfo_children()[1]:
                widget.destroy()
                break

        cars_list_frame = tk.Frame(parent_card, bg=self.colors['light'])
        cars_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_available_cars(cars_list_frame)

    def load_and_display_available_cars(self, parent_frame):
        """Загрузить и отобразить список доступных автомобилей"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/cars/available", headers=headers)
            
            if response.status_code == 200:
                cars = response.json()

                self.original_cars_data = cars

                if cars and hasattr(self, 'filter_stamp_combo'):
                    stamps = list(set(car.get('stamp', '') for car in cars if car.get('stamp')))
                    self.filter_stamp_combo['values'] = stamps

                filtered_cars = self.apply_car_filters(cars)
                
                if not filtered_cars:
                    no_cars_label = tk.Label(parent_frame, 
                                        text="По вашему запросу ничего не найдено",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 11))
                    no_cars_label.pack(pady=20)
                    return

                container, scrollable_frame, canvas, bind_scroll_to_children = self.create_scrollable_frame(parent_frame)
                container.pack(fill=tk.BOTH, expand=True)

                self.create_cars_grid(scrollable_frame, filtered_cars)

                bind_scroll_to_children(scrollable_frame)

                canvas.update_idletasks()
                canvas.configure(scrollregion=canvas.bbox("all"))
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки автомобилей")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_cars_grid(self, parent, cars):
        """Создать сетку автомобилей 4xN"""
        columns = 4
        
        for i, car in enumerate(cars):
            row = i // columns
            col = i % columns

            card_frame = self.create_car_card(car)
            card_frame.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

            parent.grid_rowconfigure(row, weight=1)
            parent.grid_columnconfigure(col, weight=1)


    def create_car_card(self, parent, car):
        """Создать карточку для отображения автомобиля"""
        card_frame = tk.Frame(parent, bg='#ffffff', relief='solid', bd=1, padx=12, pady=12)
        card_frame.config(width=280, height=220)
        card_frame.pack_propagate(False) 

        info_frame = tk.Frame(card_frame, bg='#ffffff')
        info_frame.pack(fill=tk.BOTH, expand=True)

        stamp_model = f"{car.get('stamp', '')} {car.get('model', '')}"
        if len(stamp_model) > 25:
            stamp_model = stamp_model[:25] + "..."
            
        stamp_model_label = tk.Label(info_frame, 
                                    text=f"🚗 {stamp_model}",
                                    bg='#ffffff', fg=self.colors['dark'],
                                    font=('Arial', 11, 'bold'), anchor='w')
        stamp_model_label.pack(fill=tk.X, pady=(0, 5))

        details = [
            f"📏 Пробег: {car.get('run_km', 0):,} км".replace(",", " "),
            f"💰 Стоимость: {car.get('price', 0):,} руб".replace(",", " "),
            f"🔢 VIN: {car.get('vin', '')[:12]}..." if len(car.get('vin', '')) > 12 else f"🔢 VIN: {car.get('vin', '')}"
        ]
        
        for detail in details:
            detail_label = tk.Label(info_frame, text=detail,
                                   bg='#ffffff', fg='#2c3e50',
                                   font=('Arial', 9), anchor='w')
            detail_label.pack(fill=tk.X, pady=1)

        actions_frame = tk.Frame(card_frame, bg='#ffffff')
        actions_frame.pack(fill=tk.X, pady=(8, 0))

        details_btn = ttk.Button(actions_frame, text="ℹ️ Подробнее",
                                style='Secondary.TButton',
                                command=lambda c=car: self.show_car_details(c),
                                width=12)
        details_btn.pack(side=tk.LEFT, padx=(0, 5), ipady=3, expand=True, fill=tk.X)

        buy_btn = ttk.Button(actions_frame, text="🛒 Купить",
                            style='Success.TButton',
                            command=lambda c=car: self.confirm_purchase_car(c),
                            width=12)
        buy_btn.pack(side=tk.LEFT, ipady=3, expand=True, fill=tk.X)
        
        return card_frame

    def create_cars_grid(self, parent, cars):
        """Создать сетку автомобилей 4xN"""
        columns = 4
        
        for i, car in enumerate(cars):
            row = i // columns
            col = i % columns

            card_frame = self.create_car_card(parent, car)
            card_frame.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

            parent.grid_rowconfigure(row, weight=1)
            parent.grid_columnconfigure(col, weight=1)


    def show_car_details(self, car):
        """Показать детальную информацию об автомобиле"""
        details_window = tk.Toplevel(self.root)
        details_window.title("Детали автомобиля")
        details_window.geometry("500x600")
        details_window.configure(bg=self.colors['background'])
        details_window.resizable(False, False)

        details_window.update_idletasks()
        x = (self.root.winfo_x() + (self.root.winfo_width() // 2)) - (500 // 2)
        y = (self.root.winfo_y() + (self.root.winfo_height() // 2)) - (400 // 2)
        details_window.geometry(f"500x600+{x}+{y}")
        
        details_window.transient(self.root)
        details_window.grab_set()
        
        header_frame = ttk.Frame(details_window, padding="10")
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        header_label = ttk.Label(header_frame, text="Детали автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        card = self.create_card_frame(details_window)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        info_frame = tk.Frame(card, bg=self.colors['light'])
        info_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)

        details = [
            ("🚗 Марка:", car.get('stamp', 'Не указана')),
            ("🚙 Модель:", car.get('model', 'Не указана')),
            ("📏 Пробег:", f"{car.get('run_km', 0):,} км".replace(",", " ")),
            ("💰 Цена:", f"{car.get('price', 0):,} руб".replace(",", " ")),
            ("🔢 VIN:", car.get('vin', 'Не указан')),
            ("📋 ID:", str(car.get('id', 'Не указан')))
        ]
        
        for label, value in details:
            field_frame = tk.Frame(info_frame, bg=self.colors['light'])
            field_frame.pack(fill=tk.X, pady=8)
            
            lbl = tk.Label(field_frame, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 11, 'bold'),
                        width=10, anchor='w')
            lbl.pack(side=tk.LEFT)
            
            value_lbl = tk.Label(field_frame, text=value, bg=self.colors['light'], 
                            fg='#2c3e50', font=('Arial', 11),
                            anchor='w')
            value_lbl.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))

        description = car.get('description')
        if description:
            separator = ttk.Separator(info_frame, orient='horizontal')
            separator.pack(fill=tk.X, pady=15)

            desc_header_frame = tk.Frame(info_frame, bg=self.colors['light'])
            desc_header_frame.pack(fill=tk.X, pady=(0, 8))
            
            desc_header = tk.Label(desc_header_frame, text="📄 Описание:", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 11, 'bold'), anchor='w')
            desc_header.pack(fill=tk.X)

            desc_frame = tk.Frame(info_frame, bg=self.colors['light'])
            desc_frame.pack(fill=tk.X, pady=(0, 5))
            
            desc_text = tk.Text(desc_frame, 
                            bg=self.colors['light'], fg='#2c3e50',
                            font=('Arial', 10), 
                            wrap=tk.WORD, 
                            height=4,
                            relief='flat',
                            padx=5, pady=5)
            desc_text.insert('1.0', description)
            desc_text.config(state='disabled')
            desc_text.pack(fill=tk.X)

            details_window.geometry("550x650")

        buy_btn_frame = tk.Frame(card, bg=self.colors['light'])
        buy_btn_frame.pack(fill=tk.X, pady=10)
        
        buy_btn = ttk.Button(buy_btn_frame, text="🛒 Купить этот автомобиль",
                            style='Success.TButton',
                            command=lambda: [details_window.destroy(), self.confirm_purchase_car(car)])
        buy_btn.pack(fill=tk.X, pady=6, ipady=8)

        close_btn = ttk.Button(buy_btn_frame, text="❌ Закрыть",
                            style='Secondary.TButton',
                            command=details_window.destroy)
        close_btn.pack(fill=tk.X, pady=6, ipady=6)

        details_window.focus_set()

    def confirm_purchase_car(self, car):
        """Подтверждение покупки автомобиля"""
        result = messagebox.askyesno(
            "Подтверждение покупки", 
            f"Вы уверены, что хотите купить этот автомобиль?\n\n"
            f"Марка: {car.get('stamp', '')}\n"
            f"Модель: {car.get('model', '')}\n"
            f"Цена: {car.get('price', 0):,} руб\n"
            f"VIN: {car.get('vin', '')}",
            icon='question'
        )
        
        if result:
            self.perform_purchase_car(car['id'])

    def perform_purchase_car(self, car_id):
        """Выполнить покупку автомобиля"""
        try:
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/users/cars/buy?car_id={car_id}", 
                                   headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Автомобиль успешно куплен!")
                self.show_available_cars()
            else:
                error_msg = response.json().get("detail", "Ошибка покупки автомобиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")


    def show_my_anketi(self): 
        """Показать раздел моих анкет"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Мои анкеты", 
                               style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                             style='Secondary.TButton',
                             command=self.show_main_menu,
                             width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        anketa_header = tk.Label(card, text="📝 Управление анкетами", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        anketa_header.pack(fill=tk.X, pady=(0, 15))

        create_btn_frame = tk.Frame(card, bg=self.colors['light'])
        create_btn_frame.pack(fill=tk.X, pady=10)
        
        create_btn = ttk.Button(create_btn_frame, text="➕ Создать новую анкету",
                               style='Success.TButton',
                               command=self.show_create_anketa_form,
                               width=30)
        create_btn.pack(ipady=10, anchor='center')

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        anketa_list_frame = tk.Frame(card, bg=self.colors['light'])
        anketa_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        list_header = tk.Label(anketa_list_frame, text="Мои текущие анкеты:",
                              bg=self.colors['light'], fg=self.colors['dark'],
                              font=('Arial', 12, 'bold'), anchor='w')
        list_header.pack(fill=tk.X, pady=(0, 10))

        self.load_and_display_anketi(anketa_list_frame)


    def load_and_display_anketi(self, parent_frame):
        """Загрузить и отобразить список анкет со скроллбаром"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/anketi/", headers=headers)
            
            if response.status_code == 200:
                anketi = response.json()
                
                if not anketi or (isinstance(anketi, dict) and anketi.get("message") == "Анкеты отсутствуют"):

                    no_anketi_label = tk.Label(parent_frame, 
                                            text="У вас пока нет созданных анкет",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_anketi_label.pack(pady=20)
                    return

                container, scrollable_frame, canvas, bind_scroll_to_children = self.create_scrollable_frame(parent_frame)
                container.pack(fill=tk.BOTH, expand=True)

                self.create_anketa_grid(scrollable_frame, anketi)
                
                bind_scroll_to_children(scrollable_frame)

                canvas.update_idletasks()
                canvas.configure(scrollregion=canvas.bbox("all"))
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки анкет")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_anketa_card(self, parent, anketa):
        """Создать карточку для отображения анкеты"""
        card_frame = tk.Frame(parent, bg='#ffffff', relief='solid', bd=1, padx=12, pady=12)
        card_frame.config(width=280, height=200)
        card_frame.pack_propagate(False) 

        info_frame = tk.Frame(card_frame, bg='#ffffff')
        info_frame.pack(fill=tk.BOTH, expand=True)

        stamp_model = f"{anketa.get('stamp', '')} {anketa.get('model_car', '')}"
        if len(stamp_model) > 25:
            stamp_model = stamp_model[:25] + "..."
            
        stamp_model_label = tk.Label(info_frame, 
                                    text=f"🚗 {stamp_model}",
                                    bg='#ffffff', fg=self.colors['dark'],
                                    font=('Arial', 11, 'bold'), anchor='w')
        stamp_model_label.pack(fill=tk.X, pady=(0, 5))

        vin = anketa.get('vin', '')
        display_vin = f"🔢 VIN: {vin[:12]}..." if len(vin) > 12 else f"🔢 VIN: {vin}"
        
        details = [
            f"📏 Пробег: {anketa.get('run', 0):,} км".replace(",", " "),
            f"💰 Стоимость: {anketa.get('price', 0):,} руб".replace(",", " "),
            display_vin
        ]
        
        for detail in details:
            detail_label = tk.Label(info_frame, text=detail,
                                   bg='#ffffff', fg='#2c3e50',
                                   font=('Arial', 9), anchor='w')
            detail_label.pack(fill=tk.X, pady=1)

        description = anketa.get('description', '')
        if description:
            if len(description) > 60:
                description = description[:60] + "..."
            
            desc_label = tk.Label(info_frame, text=f"📄 {description}",
                                 bg='#ffffff', fg='#7f8c8d',
                                 font=('Arial', 8), anchor='w', justify=tk.LEFT, wraplength=240)
            desc_label.pack(fill=tk.X, pady=(5, 0))

        actions_frame = tk.Frame(card_frame, bg='#ffffff')
        actions_frame.pack(fill=tk.X, pady=(8, 0))

        edit_btn = ttk.Button(actions_frame, text="✏️ Изменить",
                             style='Secondary.TButton',
                             command=lambda a=anketa: self.show_edit_anketa_form(a),
                             width=8)
        edit_btn.pack(side=tk.LEFT, padx=(0, 5), ipady=3, expand=True, fill=tk.X)

        delete_btn = ttk.Button(actions_frame, text="🗑️ Удалить",
                               style='Danger.TButton',
                               command=lambda a=anketa: self.confirm_delete_anketa(a),
                               width=8)
        delete_btn.pack(side=tk.LEFT, ipady=3, expand=True, fill=tk.X)
        
        return card_frame

    def create_anketa_grid(self, parent, anketi):
        columns = 4
        
        for i, anketa in enumerate(anketi):
            row = i // columns
            col = i % columns

            card_frame = self.create_anketa_card(parent, anketa)
            card_frame.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

            parent.grid_rowconfigure(row, weight=1)
            parent.grid_columnconfigure(col, weight=1)

    def show_create_anketa_form(self):
        """Создание новой анкеты"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Создание новой анкеты", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к списку анкет",
                            style='Secondary.TButton',
                            command=self.show_my_anketi,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')
        
        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        fields = [
            ("Марка автомобиля", "create_stamp", "combobox"),
            ("Модель автомобиля", "create_model_car", "combobox"),
            ("Пробег (км)", "create_run", "entry"),
            ("Цена (руб)", "create_price", "entry"),
            ("VIN номер", "create_vin", "entry"),
            ("Описание (необязательно)", "create_description", "text")
        ]
        
        self.create_anketa_entries = {}

        self.load_car_data()
        
        for i, (label, field_name, field_type) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            if field_type == "combobox":
                if field_name == "create_stamp":
                    stamps = [stamp['stamp'] for stamp in self.car_stamps] if hasattr(self, 'car_stamps') and self.car_stamps else []
                    entry = ttk.Combobox(field_container, font=('Arial', 11), values=stamps)
                else:
                    models = [model['model_car'] for model in self.car_models] if hasattr(self, 'car_models') and self.car_models else []
                    entry = ttk.Combobox(field_container, font=('Arial', 11), values=models)
                entry.pack(fill=tk.X, pady=2, ipady=6)
                
            elif field_type == "text":
                entry = tk.Text(field_container, font=('Arial', 11), height=4, wrap=tk.WORD)
                entry.pack(fill=tk.X, pady=2)
            else:
                entry = ttk.Entry(field_container, font=('Arial', 11))
                if "run" in field_name or "price" in field_name:
                    entry.config(validate="key", validatecommand=(self.root.register(self.validate_number), '%P'))
                entry.pack(fill=tk.X, pady=2, ipady=6)
            
            self.create_anketa_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Создать анкету", 
                            style='Success.TButton', 
                            command=self.perform_create_anketa)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_my_anketi)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def show_edit_anketa_form(self, anketa):
        """Редактирования анкеты"""
        self.clear_window()
        self.current_editing_anketa = anketa
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование анкеты", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к списку анкет",
                            style='Secondary.TButton',
                            command=self.show_my_anketi,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_car_data()

        fields = [
            ("Марка автомобиля", "edit_stamp", anketa.get('stamp', ''), "combobox"),
            ("Модель автомобиля", "edit_model_car", anketa.get('model_car', ''), "combobox"),
            ("Пробег (км)", "edit_run", str(anketa.get('run', 0)), "entry"),
            ("Цена (руб)", "edit_price", str(anketa.get('price', 0)), "entry"),
            ("VIN номер", "edit_vin", anketa.get('vin', ''), "entry"),
            ("Описание (необязательно)", "edit_description", anketa.get('description', ''), "text")
        ]
        
        self.edit_anketa_entries = {}
        
        for i, (label, field_name, current_value, field_type) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            if field_type == "combobox":
                if field_name == "edit_stamp":
                    stamps = [stamp['stamp'] for stamp in self.car_stamps] if hasattr(self, 'car_stamps') and self.car_stamps else []
                    entry = ttk.Combobox(field_container, font=('Arial', 11), values=stamps)
                else:
                    models = [model['model_car'] for model in self.car_models] if hasattr(self, 'car_models') and self.car_models else []
                    entry = ttk.Combobox(field_container, font=('Arial', 11), values=models)
                entry.set(current_value)
                entry.pack(fill=tk.X, pady=2, ipady=6)
                
            elif field_type == "text":
                entry = tk.Text(field_container, font=('Arial', 11), height=4, wrap=tk.WORD)
                entry.insert('1.0', current_value)
                entry.pack(fill=tk.X, pady=2)
            else:
                entry = ttk.Entry(field_container, font=('Arial', 11))
                entry.insert(0, current_value)
                if "run" in field_name or "price" in field_name:
                    entry.config(validate="key", validatecommand=(self.root.register(self.validate_number), '%P'))
                entry.pack(fill=tk.X, pady=2, ipady=6)
            
            self.edit_anketa_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                            style='Success.TButton', 
                            command=self.perform_edit_anketa)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_my_anketi)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def load_car_data(self):
        """Загрузить марки и модели автомобилей"""
        try:
            headers = {"token": self.auth_token}

            response = requests.get(f"{API_BASE_URL}/admin/stamps/", headers=headers)
            if response.status_code == 200:
                self.car_stamps = response.json()
            else:
                self.car_stamps = []
                print(f"Ошибка загрузки марок: {response.status_code}")

            response = requests.get(f"{API_BASE_URL}/admin/models/", headers=headers)
            if response.status_code == 200:
                self.car_models = response.json()
            else:
                self.car_models = []
                print(f"Ошибка загрузки моделей: {response.status_code}")
                
        except requests.exceptions.RequestException as e:
            self.car_stamps = []
            self.car_models = []
            print(f"Ошибка подключения: {str(e)}")


    def validate_number(self, value):
        """Валидация числовых полей"""
        if value == "":
            return True
        try:
            int(value)
            return True
        except ValueError:
            return False

    def perform_create_anketa(self):
        """Создание новой анкету"""
        try:
            data = {
                "stamp": self.get_entry_value("create_stamp"),
                "model_car": self.get_entry_value("create_model_car"),
                "run": int(self.get_entry_value("create_run")),
                "price": int(self.get_entry_value("create_price")),
                "vin": self.get_entry_value("create_vin"),
                "description": self.get_text_value("create_description")
            }

            required_fields = ["stamp", "model_car", "run", "price", "vin"]
            for field in required_fields:
                if not data[field]:
                    messagebox.showerror("Ошибка", f"Заполните поле: {field}")
                    return
            
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/users/anketa/create", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Анкета успешно создана!")
                self.show_my_anketi()
            else:
                error_msg = response.json().get("detail", "Ошибка создания анкеты")
                messagebox.showerror("Ошибка", error_msg)
                
        except ValueError:
            messagebox.showerror("Ошибка", "Пробег и цена должны быть числами")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def perform_edit_anketa(self):
        """Изменение анкеты"""
        try:
            data = {}

            fields_mapping = {
                "edit_stamp": "stamp",
                "edit_model_car": "model_car", 
                "edit_run": "run",
                "edit_price": "price",
                "edit_vin": "vin",
                "edit_description": "description"
            }
            
            for entry_name, field_name in fields_mapping.items():
                if entry_name in ["edit_run", "edit_price"]:
                    value = self.get_entry_value(entry_name)
                    if value and value != str(self.current_editing_anketa.get(field_name, '')):
                        data[field_name] = int(value)
                elif entry_name == "edit_description":
                    value = self.get_text_value(entry_name)
                    if value != self.current_editing_anketa.get(field_name, ''):
                        data[field_name] = value
                else:
                    value = self.get_entry_value(entry_name)
                    if value != self.current_editing_anketa.get(field_name, ''):
                        data[field_name] = value
            
            if not data:
                messagebox.showinfo("Информация", "Нет изменений для сохранения")
                return
            
            headers = {"token": self.auth_token}
            anketa_id = self.current_editing_anketa['id']
            response = requests.put(f"{API_BASE_URL}/users/anketa/update?anketa_id={anketa_id}", 
                                  json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Анкета успешно обновлена!")
                self.show_my_anketi()
            else:
                error_msg = response.json().get("detail", "Ошибка обновления анкеты")
                messagebox.showerror("Ошибка", error_msg)
                
        except ValueError:
            messagebox.showerror("Ошибка", "Пробег и цена должны быть числами")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_delete_anketa(self, anketa):
        """Подтверждение удаления анкеты"""
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить анкету?\n\n"
            f"Марка: {anketa.get('stamp', '')}\n"
            f"Модель: {anketa.get('model_car', '')}\n"
            f"VIN: {anketa.get('vin', '')}",
            icon='warning'
        )
        
        if result:
            self.perform_delete_anketa(anketa['id'])

    def perform_delete_anketa(self, anketa_id):
        """Удалить анкету"""
        try:
            headers = {"token": self.auth_token}
            response = requests.delete(f"{API_BASE_URL}/users/anketa/delete?anketa_id={anketa_id}", 
                                     headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Анкета успешно удалена!")
                self.show_my_anketi()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления анкеты")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def get_entry_value(self, entry_name):
        """Получить значение из Entry"""
        entry = self.create_anketa_entries.get(entry_name) or self.edit_anketa_entries.get(entry_name)
        if entry and hasattr(entry, 'get'):
            return entry.get().strip()
        return ""

    def get_text_value(self, text_name):
        """Получить значение из Text"""
        text_widget = self.create_anketa_entries.get(text_name) or self.edit_anketa_entries.get(text_name)
        if text_widget and hasattr(text_widget, 'get'):
            return text_widget.get('1.0', 'end-1c').strip()
        return ""


    def show_admin_cars_management(self):
        """Управление автомобилями для администратора"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Управление автомобилями", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        management_header = tk.Label(card, text="🚗 Управление автомобилями", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        management_header.pack(fill=tk.X, pady=(0, 15))

        add_btn_frame = tk.Frame(card, bg=self.colors['light'])
        add_btn_frame.pack(fill=tk.X, pady=10)
        
        add_btn = ttk.Button(add_btn_frame, text="➕ Добавить автомобиль",
                            style='Success.TButton',
                            command=self.show_add_car_form,
                            width=30)
        add_btn.pack(ipady=10, anchor='center')
        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        self.current_filter_callback = lambda: self.refresh_admin_cars_list(card)
        filter_frame = self.show_filter_sort_options(card, self.current_filter_callback)
        
        cars_list_frame = tk.Frame(card, bg=self.colors['light'])
        cars_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_all_cars(cars_list_frame)

    def refresh_admin_cars_list(self, parent_card):
        """Обновить список автомобилей в админке"""
        for widget in parent_card.winfo_children():
            if isinstance(widget, tk.Frame) and widget not in [parent_card.winfo_children()[1], 
                                                            parent_card.winfo_children()[2],
                                                            parent_card.winfo_children()[3]]:
                widget.destroy()
                break

        cars_list_frame = tk.Frame(parent_card, bg=self.colors['light'])
        cars_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_all_cars(cars_list_frame)

    def show_admin_anketi_management(self):
        """Управление анкетами для администратора"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Анкеты пользователей", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        anketa_header = tk.Label(card, text="📝 Анкеты пользователей", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        anketa_header.pack(fill=tk.X, pady=(0, 15))

        anketa_list_frame = tk.Frame(card, bg=self.colors['light'])
        anketa_list_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_admin_anketi_like_user(anketa_list_frame)

    def load_and_display_admin_anketi_like_user(self, parent_frame):
        """Загрузить и отобразить анкеты в стиле пользовательского интерфейса"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/anketi/", headers=headers)
            
            if response.status_code == 200:
                anketi = response.json()
                
                if not anketi:
                    no_anketi_label = tk.Label(parent_frame, 
                                            text="Нет анкет от пользователей",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_anketi_label.pack(pady=20)
                    return

                container, scrollable_frame, canvas, bind_scroll_to_children = self.create_scrollable_frame(parent_frame)
                container.pack(fill=tk.BOTH, expand=True)

                self.create_admin_anketa_grid_like_user(scrollable_frame, anketi)

                bind_scroll_to_children(scrollable_frame)

                canvas.update_idletasks()
                canvas.configure(scrollregion=canvas.bbox("all"))
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки анкет")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_admin_anketa_grid_like_user(self, parent, anketi):
        """Создать сетку анкет в стиле пользовательского интерфейса"""
        columns = 4
        
        for i, anketa in enumerate(anketi):
            row = i // columns
            col = i % columns

            card_frame = self.create_admin_anketa_card_like_user(parent, anketa)
            card_frame.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

            parent.grid_rowconfigure(row, weight=1)
            parent.grid_columnconfigure(col, weight=1)

    def create_admin_anketa_card_like_user(self, parent, anketa):
        """Создать карточку анкеты для администратора в стиле пользовательского интерфейса"""
        card_frame = tk.Frame(parent, bg='#ffffff', relief='solid', bd=1, padx=12, pady=12)
        card_frame.config(width=280, height=300)
        card_frame.pack_propagate(False)

        info_frame = tk.Frame(card_frame, bg='#ffffff')
        info_frame.pack(fill=tk.BOTH, expand=True)

        user_info = f"👤 {anketa.get('user_name', '')}"
        user_label = tk.Label(info_frame, 
                            text=user_info,
                            bg='#ffffff', fg=self.colors['dark'],
                            font=('Arial', 10, 'bold'), anchor='w')
        user_label.pack(fill=tk.X, pady=(0, 3))

        stamp_model = f"{anketa.get('stamp', '')} {anketa.get('model_car', '')}"
        if len(stamp_model) > 25:
            stamp_model = stamp_model[:25] + "..."
            
        stamp_model_label = tk.Label(info_frame, 
                                    text=f"🚗 {stamp_model}",
                                    bg='#ffffff', fg=self.colors['dark'],
                                    font=('Arial', 11, 'bold'), anchor='w')
        stamp_model_label.pack(fill=tk.X, pady=(0, 5))
    
        vin = anketa.get('vin', '')
        display_vin = f"🔢 VIN: {vin[:12]}..." if len(vin) > 12 else f"🔢 VIN: {vin}"
        
        details = [
            f"📏 Пробег: {anketa.get('run', 0):,} км".replace(",", " "),
            f"💰 Стоимость: {anketa.get('price', 0):,} руб".replace(",", " "),
            display_vin
        ]
        
        for detail in details:
            detail_label = tk.Label(info_frame, text=detail,
                                bg='#ffffff', fg='#2c3e50',
                                font=('Arial', 9), anchor='w')
            detail_label.pack(fill=tk.X, pady=1)
        
        description = anketa.get('description', '')
        if description:
            if len(description) > 60:
                description = description[:60] + "..."
            
            desc_label = tk.Label(info_frame, text=f"📄 {description}",
                                bg='#ffffff', fg='#7f8c8d',
                                font=('Arial', 8), anchor='w', justify=tk.LEFT, wraplength=240)
            desc_label.pack(fill=tk.X, pady=(5, 0))

        actions_frame = tk.Frame(card_frame, bg='#ffffff')
        actions_frame.pack(fill=tk.X, pady=(8, 0))

        accept_btn = ttk.Button(actions_frame, text="✅ Принять анкету",
                            style='Success.TButton',
                            command=lambda a=anketa: self.confirm_accept_anketa(a))
        accept_btn.pack(fill=tk.X, ipady=3)
        
        return card_frame
    
    def show_user_management(self):
        """Управление пользователями для администратора с Treeview"""
        self.clear_window()
        self.selected_user = None
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Управление пользователями", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        management_header = tk.Label(card, text="👥 Управление пользователями", 
                                    bg=self.colors['light'], fg=self.colors['dark'],
                                    font=('Arial', 14, 'bold'), anchor='w')
        management_header.pack(fill=tk.X, pady=(0, 15))

        actions_frame = tk.Frame(card, bg=self.colors['light'])
        actions_frame.pack(fill=tk.X, pady=10)

        buttons_container = tk.Frame(actions_frame, bg=self.colors['light'])
        buttons_container.pack(fill=tk.X)
        
        create_btn = ttk.Button(buttons_container, text="➕ Создать пользователя",
                            style='Success.TButton',
                            command=self.show_create_user_form,
                            width=20)
        create_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        edit_btn = ttk.Button(buttons_container, text="✏️ Редактировать",
                            style='Accent.TButton',
                            command=self.show_edit_user_form,
                            width=20)
        edit_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        delete_btn = ttk.Button(buttons_container, text="🗑️ Удалить",
                            style='Danger.TButton',
                            command=self.confirm_delete_user,
                            width=20)
        delete_btn.pack(side=tk.LEFT, ipady=8)

        self.selected_user_info = tk.Label(actions_frame, 
                                        text="Выберите пользователя из таблицы",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 10), anchor='w')
        self.selected_user_info.pack(fill=tk.X, pady=(10, 0))

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        table_frame = tk.Frame(card, bg=self.colors['light'])
        table_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        table_header = tk.Label(table_frame, text="Список пользователей:",
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 12, 'bold'), anchor='w')
        table_header.pack(fill=tk.X, pady=(0, 10))

        self.load_and_display_users_table(table_frame)

    def load_and_display_users_table(self, parent_frame):
        """Загрузить и отобразить таблицу пользователей с Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/list_users/", headers=headers)
            
            if response.status_code == 200:
                users = response.json()
                
                if not users:
                    no_users_label = tk.Label(parent_frame, 
                                            text="В системе нет других пользователей",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_users_label.pack(pady=20)
                    return

                self.create_users_treeview(parent_frame, users)
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки пользователей")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_users_treeview(self, parent, users):
        """Создать Treeview для отображения пользователей"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        h_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        self.users_tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Имя', 'Email', 'Телефон'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            xscrollcommand=h_scrollbar.set,
            height=15
        )

        v_scrollbar.config(command=self.users_tree.yview)
        h_scrollbar.config(command=self.users_tree.xview)

        self.users_tree.heading('ID', text='ID', anchor=tk.W)
        self.users_tree.heading('Имя', text='Имя', anchor=tk.W)
        self.users_tree.heading('Email', text='Email', anchor=tk.W)
        self.users_tree.heading('Телефон', text='Телефон', anchor=tk.W)

        self.users_tree.column('ID', width=60, minwidth=50)
        self.users_tree.column('Имя', width=200, minwidth=150)
        self.users_tree.column('Email', width=250, minwidth=200)
        self.users_tree.column('Телефон', width=150, minwidth=120)

        self.users_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for user in sorted(users, key=lambda x: x['id']):
            self.users_tree.insert(
                '', 
                tk.END, 
                values=(
                    user['id'],
                    user['name'],
                    user['email'],
                    user['phone']
                )
            )

        self.users_tree.bind('<<TreeviewSelect>>', self.on_user_select)

        self.configure_treeview_style()

    def configure_treeview_style(self):
        """Настроить стиль для Treeview"""
        style = ttk.Style()
        style.theme_use('clam')

        style.configure(
            'Treeview',
            background='white',
            foreground='#2c3e50',
            fieldbackground='white',
            borderwidth=1,
            relief='solid'
        )
        
        style.configure(
            'Treeview.Heading',
            background=self.colors['dark'],
            foreground='white',
            relief='flat',
            font=('Arial', 10, 'bold')
        )

        style.map(
            'Treeview',
            background=[('selected', self.colors['primary'])],
            foreground=[('selected', 'white')]
        )

    def on_user_select(self, event):
        """Обработчик выбора пользователя в Treeview"""
        selected_items = self.users_tree.selection()
        if selected_items:
            item = selected_items[0]
            user_data = self.users_tree.item(item, 'values')
            
            self.selected_user = {
                'id': int(user_data[0]),
                'name': user_data[1],
                'email': user_data[2],
                'phone': user_data[3]
            }
            
            self.selected_user_info.config(
                text=f"Выбран: {self.selected_user['name']} (ID: {self.selected_user['id']}, Email: {self.selected_user['email']})",
                fg=self.colors['dark']
            )
        else:
            self.selected_user = None
            self.selected_user_info.config(
                text="Выберите пользователя из таблицы",
                fg='#7f8c8d'
            )

    def show_create_user_form(self):
        """Показать форму создания пользователя"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Создание пользователя", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению пользователями",
                            style='Secondary.TButton',
                            command=self.show_user_management,
                            width=37)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        fields = [
            ("Полное имя", "create_user_name"),
            ("Email", "create_user_email"),
            ("Телефон", "create_user_phone"),
            ("Пароль", "create_user_password")
        ]
        
        self.create_user_entries = {}
        
        for i, (label, field_name) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_container, font=('Arial', 11))
            if "password" in field_name:
                entry.config(show="•")
            entry.pack(fill=tk.X, pady=2, ipady=6)
            self.create_user_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Создать пользователя", 
                            style='Success.TButton', 
                            command=self.perform_create_user)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_user_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def show_edit_user_form(self):
        """Показать форму редактирования пользователя"""
        if not self.selected_user:
            messagebox.showwarning("Внимание", "Сначала выберите пользователя из таблицы")
            return
        
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование пользователя", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению пользователями",
                            style='Secondary.TButton',
                            command=self.show_user_management,
                            width=37)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        info_frame = tk.Frame(form_frame, bg=self.colors['light'])
        info_frame.pack(fill=tk.X, pady=(0, 20))
        
        info_text = f"Редактирование пользователя:\nID: {self.selected_user['id']} • Email: {self.selected_user['email']}"
        info_label = tk.Label(info_frame, text=info_text,
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 11, 'bold'), anchor='w', justify=tk.LEFT)
        info_label.pack(fill=tk.X)

        fields = [
            ("Полное имя", "edit_user_name", self.selected_user.get('name', '')),
            ("Телефон", "edit_user_phone", self.selected_user.get('phone', ''))
        ]
        
        self.edit_user_entries = {}
        
        for i, (label, field_name, current_value) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=12)

            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))

            entry = ttk.Entry(field_container, font=('Arial', 11))
            entry.insert(0, current_value)
            entry.pack(fill=tk.X, pady=2, ipady=6)
            self.edit_user_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                            style='Success.TButton', 
                            command=self.perform_edit_user)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_user_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def perform_create_user(self):
        """Создать нового пользователя"""
        try:
            data = {
                "full_name": self.create_user_entries["create_user_name"].get().strip(),
                "email": self.create_user_entries["create_user_email"].get().strip(),
                "phone": self.create_user_entries["create_user_phone"].get().strip(),
                "password": self.create_user_entries["create_user_password"].get()
            }

            for field, value in data.items():
                if not value:
                    messagebox.showerror("Ошибка", "Заполните все поля")
                    return
            
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/admin/users/", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Пользователь успешно создан!")
                self.show_user_management()
            else:
                error_msg = response.json().get("detail", "Ошибка создания пользователя")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def perform_edit_user(self):
        """Редактировать пользователя"""
        try:
            data = {}

            if self.edit_user_entries["edit_user_name"].get().strip():
                data["full_name"] = self.edit_user_entries["edit_user_name"].get().strip()
            
            if self.edit_user_entries["edit_user_phone"].get().strip():
                data["phone"] = self.edit_user_entries["edit_user_phone"].get().strip()
            
            if not data:
                messagebox.showinfo("Информация", "Нет изменений для сохранения")
                return
            
            headers = {"token": self.auth_token}
            user_id = self.selected_user['id']
            response = requests.put(f"{API_BASE_URL}/admin/users/{user_id}", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Пользователь успешно обновлен!")
                self.show_user_management()
            else:
                error_msg = response.json().get("detail", "Ошибка обновления пользователя")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_delete_user(self):
        """Подтверждение удаления пользователя"""
        if not self.selected_user:
            messagebox.showwarning("Внимание", "Сначала выберите пользователя из таблицы")
            return
        
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить пользователя?\n\n"
            f"ID: {self.selected_user['id']}\n"
            f"Имя: {self.selected_user['name']}\n"
            f"Email: {self.selected_user['email']}\n"
            f"Телефон: {self.selected_user['phone']}",
            icon='warning'
        )
        
        if result:
            self.perform_delete_user()

    def perform_delete_user(self):
        """Удалить пользователя"""
        try:
            headers = {"token": self.auth_token}
            user_id = self.selected_user['id']
            response = requests.delete(f"{API_BASE_URL}/admin/users/delete_profile/?user_id={user_id}", 
                                    headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Пользователь успешно удален!")
                self.selected_user = None
                self.show_user_management()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления пользователя")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def show_my_purchases(self):
        """Показать мои покупки с Treeview"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Мои покупки", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        purchases_header = tk.Label(header_content, text="🛒 История моих покупок", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        purchases_header.pack(side=tk.LEFT, fill=tk.X, expand=True)

        purchases_frame = tk.Frame(card, bg=self.colors['light'])
        purchases_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_my_purchases_treeview(purchases_frame)

    def load_and_display_my_purchases_treeview(self, parent_frame):
        """Загрузить и отобразить покупки пользователя в Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/users/my_purchases", headers=headers)
            
            purchases = []
            
            if response.status_code == 200:
                purchases = response.json()
            elif response.status_code == 404:
                pass
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки покупок")
                messagebox.showerror("Ошибка", error_msg)
                return
            
            if not purchases:
                self.show_no_purchases_message(parent_frame)
            else:
                self.create_purchases_treeview(parent_frame, purchases)
                
        except requests.exceptions.RequestException as e:
            self.show_no_purchases_message(parent_frame)
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def show_no_purchases_message(self, parent_frame):
        """Показать сообщение об отсутствии покупок"""
        no_purchases_label = tk.Label(parent_frame, 
                                    text="У вас пока нет покупок",
                                    bg=self.colors['light'], fg='#7f8c8d',
                                    font=('Arial', 11))
        no_purchases_label.pack(pady=20)

    def create_purchases_treeview(self, parent, purchases):
        """Создать Treeview для отображения покупок"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Марка', 'Модель', 'VIN', 'Цена', 'Дата покупки'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            height=10  
        )

        v_scrollbar.config(command=tree.yview)

        columns = {
            'ID': {'text': 'ID', 'width': 60, 'anchor': tk.CENTER},
            'Марка': {'text': 'Марка', 'width': 120, 'anchor': tk.W},
            'Модель': {'text': 'Модель', 'width': 120, 'anchor': tk.W},
            'VIN': {'text': 'VIN', 'width': 150, 'anchor': tk.W},
            'Цена': {'text': 'Цена', 'width': 120, 'anchor': tk.E},
            'Дата покупки': {'text': 'Дата покупки', 'width': 120, 'anchor': tk.CENTER}
        }

        for col, settings in columns.items():
            tree.heading(col, text=settings['text'], anchor=settings['anchor'])
            tree.column(col, width=settings['width'], minwidth=50, anchor=settings['anchor'])

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for purchase in sorted(purchases, key=lambda x: x['id']):
            car_info = purchase.get('car', {})
            price = purchase.get('price', 0)
            date_buy = purchase.get('date_buy', '')[:10]
            
            tree.insert(
                '', 
                tk.END, 
                values=(
                    purchase['id'],
                    car_info.get('stamp', ''),
                    car_info.get('model', ''),
                    car_info.get('vin', ''),
                    f"{price:,} руб".replace(",", " "),
                    date_buy
                )
            )


        self.configure_treeview_style()

    def show_admin_purchases(self):
        """История покупок для администратора с Treeview"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="История покупок", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        purchases_header = tk.Label(card, text="🛒 История покупок пользователей", 
                                bg=self.colors['light'], fg=self.colors['dark'],
                                font=('Arial', 14, 'bold'), anchor='w')
        purchases_header.pack(fill=tk.X, pady=(0, 15))

        purchases_frame = tk.Frame(card, bg=self.colors['light'])
        purchases_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_all_purchases_treeview(purchases_frame)

    def load_and_display_all_purchases_treeview(self, parent_frame):
        """Загрузить и отобразить все покупки в Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/shopping/", headers=headers)
            
            if response.status_code == 200:
                purchases = response.json()
                
                if not purchases:
                    no_data_label = tk.Label(parent_frame, 
                                        text="Нет данных о покупках",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 11))
                    no_data_label.pack(pady=20)
                    return
                
                self.create_admin_purchases_treeview(parent_frame, purchases)
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки данных о покупках")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_admin_purchases_treeview(self, parent, purchases):
        """Создать Treeview для отображения всех покупок (админ)"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Покупатель', 'Email', 'Марка', 'Модель', 'VIN', 'Цена', 'Дата'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            height=10
        )

        v_scrollbar.config(command=tree.yview)

        columns = {
            'ID': {'text': 'ID', 'width': 60, 'anchor': tk.CENTER},
            'Покупатель': {'text': 'Покупатель', 'width': 150, 'anchor': tk.W},
            'Email': {'text': 'Email', 'width': 200, 'anchor': tk.W},
            'Марка': {'text': 'Марка', 'width': 120, 'anchor': tk.W},
            'Модель': {'text': 'Модель', 'width': 120, 'anchor': tk.W},
            'VIN': {'text': 'VIN', 'width': 150, 'anchor': tk.W},
            'Цена': {'text': 'Цена', 'width': 120, 'anchor': tk.E},
            'Дата': {'text': 'Дата покупки', 'width': 120, 'anchor': tk.CENTER}
        }

        for col, settings in columns.items():
            tree.heading(col, text=settings['text'], anchor=settings['anchor'])
            tree.column(col, width=settings['width'], minwidth=50, anchor=settings['anchor'])

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for purchase in sorted(purchases, key=lambda x: x['id']):
            car_info = purchase.get('car', {})
            buyer_info = purchase.get('buyer', {})
            price = purchase.get('price', 0)
            date_buy = purchase.get('date_buy', '')[:10]
            
            tree.insert(
                '', 
                tk.END, 
                values=(
                    purchase['id'],
                    buyer_info.get('name', ''),
                    buyer_info.get('email', ''),
                    car_info.get('stamp', ''),
                    car_info.get('model', ''),
                    car_info.get('vin', ''),
                    f"{price:,} руб".replace(",", " "),
                    date_buy
                )
            )

        self.configure_treeview_style()

    def show_admin_sales(self):
        """История продаж для администратора с Treeview"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="История продаж", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        sales_header = tk.Label(card, text="💰 История продаж автомобилей", 
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 14, 'bold'), anchor='w')
        sales_header.pack(fill=tk.X, pady=(0, 15))

        sales_frame = tk.Frame(card, bg=self.colors['light'])
        sales_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.load_and_display_all_sales_treeview(sales_frame)

    def load_and_display_all_sales_treeview(self, parent_frame):
        """Загрузить и отобразить все продажи в Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/sales/", headers=headers)
            
            if response.status_code == 200:
                sales = response.json()
                
                if not sales:
                    no_data_label = tk.Label(parent_frame, 
                                        text="Нет данных о продажах",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 11))
                    no_data_label.pack(pady=20)
                    return
                
                self.create_admin_sales_treeview(parent_frame, sales)
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки данных о продажах")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_admin_sales_treeview(self, parent, sales):
        """Создать Treeview для отображения всех продаж (админ)"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Покупатель', 'Email', 'Марка', 'Модель', 'VIN', 'Цена', 'Дата'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            height=10
        )

        v_scrollbar.config(command=tree.yview)

        columns = {
            'ID': {'text': 'ID', 'width': 60, 'anchor': tk.CENTER},
            'Покупатель': {'text': 'Покупатель', 'width': 150, 'anchor': tk.W},
            'Email': {'text': 'Email', 'width': 200, 'anchor': tk.W},
            'Марка': {'text': 'Марка', 'width': 120, 'anchor': tk.W},
            'Модель': {'text': 'Модель', 'width': 120, 'anchor': tk.W},
            'VIN': {'text': 'VIN', 'width': 150, 'anchor': tk.W},
            'Цена': {'text': 'Цена', 'width': 120, 'anchor': tk.E},
            'Дата': {'text': 'Дата продажи', 'width': 120, 'anchor': tk.CENTER}
        }

        for col, settings in columns.items():
            tree.heading(col, text=settings['text'], anchor=settings['anchor'])
            tree.column(col, width=settings['width'], minwidth=50, anchor=settings['anchor'])

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for sale in sorted(sales, key=lambda x: x['id']):
            car_info = sale.get('car', {})
            buyer_info = sale.get('buyer', {})
            price = sale.get('price', 0)
            date_sale = sale.get('date_sale', '')[:10]
            
            tree.insert(
                '', 
                tk.END, 
                values=(
                    sale['id'],
                    buyer_info.get('name', ''),
                    buyer_info.get('email', ''),
                    car_info.get('stamp', ''),
                    car_info.get('model', ''),
                    car_info.get('vin', ''),
                    f"{price:,} руб".replace(",", " "),
                    date_sale
                )
            )

        self.configure_treeview_style()

    def configure_treeview_style(self):
        """Настроить стиль для Treeview"""
        style = ttk.Style()
        style.theme_use('clam')

        style.configure(
            'Treeview',
            background='white',
            foreground='#2c3e50',
            fieldbackground='white',
            borderwidth=1,
            relief='solid',
            rowheight=25
        )
        
        style.configure(
            'Treeview.Heading',
            background=self.colors['dark'],
            foreground='white',
            relief='flat',
            font=('Arial', 10, 'bold')
        )

        style.map(
            'Treeview',
            background=[('selected', self.colors['primary'])],
            foreground=[('selected', 'white')]
        )

    def show_add_car_form(self):
        """Форма добавления автомобиля для администратора"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Добавление автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению автомобилями",
                            style='Secondary.TButton',
                            command=self.show_admin_cars_management,
                            width=37)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        fields = [
            ("Марка автомобиля", "add_stamp", ),
            ("Модель автомобиля", "add_model"), 
            ("Пробег (км)", "add_run"),
            ("Цена (руб)", "add_price"),
            ("VIN номер", "add_vin"),
            ("Описание", "add_description")
        ]
        
        self.add_car_entries = {}
        
        for i, (label, field_name) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=8)
            
            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))
            
            if field_name == "add_description":
                entry = tk.Text(field_container, font=('Arial', 11), height=3, wrap=tk.WORD)
                entry.pack(fill=tk.X, pady=2)
            else:
                entry = ttk.Entry(field_container, font=('Arial', 11))
                if field_name in ["add_run", "add_price"]:
                    entry.config(validate="key", validatecommand=(self.root.register(self.validate_number), '%P'))
                entry.pack(fill=tk.X, pady=2, ipady=6)
            
            self.add_car_entries[field_name] = entry

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Добавить автомобиль", 
                            style='Success.TButton', 
                            command=self.perform_add_car)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_admin_cars_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def perform_add_car(self):
        """Добавить автомобиль"""
        try:
            data = {
                "stamp": self.add_car_entries["add_stamp"].get().strip(),
                "model_car": self.add_car_entries["add_model"].get().strip(),
                "run": int(self.add_car_entries["add_run"].get()),
                "price": int(self.add_car_entries["add_price"].get()),
                "vin": self.add_car_entries["add_vin"].get().strip(),
                "description": self.add_car_entries["add_description"].get("1.0", "end-1c").strip()
            }

            required_fields = ["stamp", "model_car", "run", "price", "vin"]
            for field in required_fields:
                if not data[field]:
                    messagebox.showerror("Ошибка", f"Заполните поле: {field}")
                    return
            
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/admin/cars/", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Автомобиль успешно добавлен!")
                self.show_admin_cars_management()
            else:
                error_msg = response.json().get("detail", "Ошибка добавления автомобиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except ValueError:
            messagebox.showerror("Ошибка", "Пробег и цена должны быть числами")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def load_and_display_all_cars(self, parent_frame):
        """Загрузить и отобразить все автомобили для администратора"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/cars/", headers=headers)
            
            if response.status_code == 200:
                cars = response.json()

                self.original_cars_data = cars

                if cars:
                    stamps = list(set(car.get('stamp', '') for car in cars if car.get('stamp')))
                    if hasattr(self, 'filter_stamp_combo'):
                        self.filter_stamp_combo['values'] = stamps

                filtered_cars = self.apply_car_filters(cars)
                
                if not filtered_cars:
                    no_cars_label = tk.Label(parent_frame, 
                                        text="По вашему запросу ничего не найдено",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 11))
                    no_cars_label.pack(pady=20)
                    return

                container, scrollable_frame, canvas, bind_scroll_to_children = self.create_scrollable_frame(parent_frame)
                container.pack(fill=tk.BOTH, expand=True)

                self.create_admin_cars_grid(scrollable_frame, filtered_cars)

                bind_scroll_to_children(scrollable_frame)

                canvas.update_idletasks()
                canvas.configure(scrollregion=canvas.bbox("all"))
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки автомобилей")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def load_and_display_all_anketi(self, parent_frame):
        """Загрузить и отобразить все анкеты для администратора"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/anketi/", headers=headers)
            
            if response.status_code == 200:
                anketi = response.json()
                
                if not anketi:
                    no_anketi_label = tk.Label(parent_frame, 
                                            text="Нет анкет от пользователей",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_anketi_label.pack(pady=20)
                    return

                container, scrollable_frame, canvas, bind_scroll_to_children = self.create_scrollable_frame(parent_frame)
                container.pack(fill=tk.BOTH, expand=True)

                self.create_admin_anketi_list(scrollable_frame, anketi)

                bind_scroll_to_children(scrollable_frame)

                canvas.update_idletasks()
                canvas.configure(scrollregion=canvas.bbox("all"))
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки анкет")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_admin_cars_grid(self, parent, cars):
        """Создать сетку автомобилей для администратора с кнопками управления"""
        columns = 4
        
        for i, car in enumerate(cars):
            row = i // columns
            col = i % columns

            card_frame = self.create_admin_car_card(parent, car)
            card_frame.grid(row=row, column=col, padx=5, pady=5, sticky="nsew")

            parent.grid_rowconfigure(row, weight=1)
            parent.grid_columnconfigure(col, weight=1)

    def create_admin_car_card(self, parent, car):
        """Создать карточку автомобиля для администратора"""
        card_frame = tk.Frame(parent, bg='#ffffff', relief='solid', bd=1, padx=10, pady=10)
        card_frame.config(width=250, height=180)
        card_frame.pack_propagate(False)

        info_frame = tk.Frame(card_frame, bg='#ffffff')
        info_frame.pack(fill=tk.BOTH, expand=True)

        stamp_model = f"{car.get('stamp', '')} {car.get('model', '')}"
        if len(stamp_model) > 20:
            stamp_model = stamp_model[:20] + "..."
            
        stamp_model_label = tk.Label(info_frame, 
                                    text=f"🚗 {stamp_model}",
                                    bg='#ffffff', fg=self.colors['dark'],
                                    font=('Arial', 10, 'bold'), anchor='w')
        stamp_model_label.pack(fill=tk.X, pady=(0, 3))

        status_label = tk.Label(info_frame, 
                            text=f"📊 Статус: {car.get('status', '')}",
                            bg='#ffffff', fg='#2c3e50',
                            font=('Arial', 8), anchor='w')
        status_label.pack(fill=tk.X, pady=1)

        details = [
            f"📏 Пробег: {car.get('run_km', 0):,} км".replace(",", " "),
            f"💰 Стоимость:  {car.get('price', 0):,} руб".replace(",", " "),
        ]
        
        for detail in details:
            detail_label = tk.Label(info_frame, text=detail,
                                bg='#ffffff', fg='#2c3e50',
                                font=('Arial', 8), anchor='w')
            detail_label.pack(fill=tk.X, pady=1)

        vin = car.get('vin', '')
        vin_display = f"🔢 VIN: {vin[:8]}..." if len(vin) > 8 else f"🔢 VIN: {vin}"
        vin_label = tk.Label(info_frame, text=vin_display,
                            bg='#ffffff', fg='#2c3e50',
                            font=('Arial', 8), anchor='w')
        vin_label.pack(fill=tk.X, pady=1)

        actions_frame = tk.Frame(card_frame, bg='#ffffff')
        actions_frame.pack(fill=tk.X, pady=(5, 0))

        edit_btn = ttk.Button(actions_frame, text="✏️Изменить",
                            style='Secondary.TButton',
                            command=lambda c=car: self.show_edit_car_form(c),
                            width=10)
        edit_btn.pack(side=tk.LEFT, padx=(0, 3), ipady=1, fill=tk.X, expand=True)

        delete_btn = ttk.Button(actions_frame, text="🗑️Удалить",
                            style='Danger.TButton',
                            command=lambda c=car: self.confirm_delete_car(c),
                            width=10)
        delete_btn.pack(side=tk.LEFT, ipady=1, fill=tk.X, expand=True)
        
        return card_frame

    def create_admin_anketi_list(self, parent, anketi):
        """Создать список анкет для администратора"""
        for anketa in anketi:
            self.create_admin_anketa_card(parent, anketa)

    def create_admin_anketa_card(self, parent, anketa):
        """Создать карточку анкеты для администратора"""
        card_frame = tk.Frame(parent, bg='#ffffff', relief='solid', bd=1, padx=15, pady=15)
        card_frame.pack(fill=tk.X, pady=8, padx=5)

        info_frame = tk.Frame(card_frame, bg='#ffffff')
        info_frame.pack(fill=tk.X)

        user_info = f"👤 {anketa.get('user_name', '')} ({anketa.get('user_phone', '')})"
        user_label = tk.Label(info_frame, text=user_info,
                            bg='#ffffff', fg=self.colors['dark'],
                            font=('Arial', 11, 'bold'), anchor='w')
        user_label.pack(fill=tk.X, pady=(0, 5))

        car_info = f"🚗 {anketa.get('stamp', '')} {anketa.get('model_car', '')}"
        car_label = tk.Label(info_frame, text=car_info,
                            bg='#ffffff', fg='#2c3e50',
                            font=('Arial', 10), anchor='w')
        car_label.pack(fill=tk.X, pady=2)

        details = [
            f"📏 Пробег: {anketa.get('run', 0):,} км".replace(",", " "),
            f"💰 Предлагаемая цена: {anketa.get('price', 0):,} руб".replace(",", " "),
            f"🔢 VIN: {anketa.get('vin', '')}"
        ]
        
        for detail in details:
            detail_label = tk.Label(info_frame, text=detail,
                                bg='#ffffff', fg='#2c3e50',
                                font=('Arial', 9), anchor='w')
            detail_label.pack(fill=tk.X, pady=1)

        description = anketa.get('description', '')
        if description:
            desc_label = tk.Label(info_frame, text=f"📄 Описание: {description}",
                                bg='#ffffff', fg='#7f8c8d',
                                font=('Arial', 8), anchor='w', justify=tk.LEFT, wraplength=500)
            desc_label.pack(fill=tk.X, pady=(5, 0))
        
        accept_btn = ttk.Button(card_frame, text="✅ Принять анкету",
                            style='Success.TButton',
                            command=lambda a=anketa: self.confirm_accept_anketa(a),
                            width=20)
        accept_btn.pack(fill=tk.X, pady=(10, 0), ipady=5)

    def show_edit_car_form(self, car):
        """Форма редактирования автомобиля"""
        self.clear_window()
        self.current_editing_car = car
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению автомобилями",
                            style='Secondary.TButton',
                            command=self.show_admin_cars_management,
                            width=37)
        back_btn.pack(ipady=8, anchor='center')
        
        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        fields = [
            ("Пробег (км)", "edit_run", str(car.get('run_km', 0))),
            ("Цена (руб)", "edit_price", str(car.get('price', 0))),
            ("Описание", "edit_description", car.get('description', ''))
        ]
        
        self.edit_car_entries = {}
        
        for i, (label, field_name, current_value) in enumerate(fields):
            field_container = tk.Frame(form_frame, bg=self.colors['light'])
            field_container.pack(fill=tk.X, pady=8)
            
            lbl = tk.Label(field_container, text=label, bg=self.colors['light'], 
                        fg=self.colors['dark'], font=('Arial', 10, 'bold'),
                        anchor='w')
            lbl.pack(fill=tk.X, pady=(0, 5))
            
            if field_name == "edit_description":
                entry = tk.Text(field_container, font=('Arial', 11), height=3, wrap=tk.WORD)
                entry.insert('1.0', current_value)
                entry.pack(fill=tk.X, pady=2)
            else:
                entry = ttk.Entry(field_container, font=('Arial', 11))
                entry.insert(0, current_value)
                if field_name in ["edit_run", "edit_price"]:
                    entry.config(validate="key", validatecommand=(self.root.register(self.validate_number), '%P'))
                entry.pack(fill=tk.X, pady=2, ipady=6)
            
            self.edit_car_entries[field_name] = entry
        
        info_frame = tk.Frame(form_frame, bg=self.colors['light'])
        info_frame.pack(fill=tk.X, pady=15)
        
        info_text = f"🚗 {car.get('stamp', '')} {car.get('model', '')}\n🔢 VIN: {car.get('vin', '')}\n📊 Статус: {car.get('status', '')}"
        info_label = tk.Label(info_frame, text=info_text,
                            bg=self.colors['light'], fg='#2c3e50',
                            font=('Arial', 11), anchor='w', justify=tk.LEFT)
        info_label.pack(fill=tk.X)

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))
        
        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                            style='Success.TButton', 
                            command=self.perform_edit_car)
        save_btn.pack(fill=tk.X, pady=6, ipady=8)
        

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_admin_cars_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=6)

    def perform_edit_car(self):
        """Редактировать автомобиль"""
        try:
            data = {}

            if self.edit_car_entries["edit_run"].get():
                data["run"] = int(self.edit_car_entries["edit_run"].get())
            
            if self.edit_car_entries["edit_price"].get():
                data["price"] = int(self.edit_car_entries["edit_price"].get())
            
            description = self.edit_car_entries["edit_description"].get("1.0", "end-1c").strip()
            if description != self.current_editing_car.get('description', ''):
                data["description"] = description
            
            if not data:
                messagebox.showinfo("Информация", "Нет изменений для сохранения")
                return
            
            headers = {"token": self.auth_token}
            car_id = self.current_editing_car['id']
            response = requests.put(f"{API_BASE_URL}/admin/cars/{car_id}", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Автомобиль успешно обновлен!")
                self.show_admin_cars_management()
            else:
                error_msg = response.json().get("detail", "Ошибка обновления автомобиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except ValueError:
            messagebox.showerror("Ошибка", "Пробег и цена должны быть числами")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")


    def confirm_delete_car(self, car):
        """Подтверждение удаления автомобиля"""
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить автомобиль?\n\n"
            f"Марка: {car.get('stamp', '')}\n"
            f"Модель: {car.get('model', '')}\n"
            f"VIN: {car.get('vin', '')}",
            icon='warning'
        )
        
        if result:
            self.perform_delete_car(car['id'])

    def perform_delete_car(self, car_id):
        """Удалить автомобиль"""
        try:
            headers = {"token": self.auth_token}
            response = requests.delete(f"{API_BASE_URL}/admin/cars/{car_id}", headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Автомобиль успешно удален!")
                self.show_admin_cars_management()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления автомобиля")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_accept_anketa(self, anketa):
        """Подтверждение принятия анкеты"""
        result = messagebox.askyesno(
            "Подтверждение принятия", 
            f"Вы уверены, что хотите принять эту анкету?\n\n"
            f"Пользователь: {anketa.get('user_name', '')}\n"
            f"Автомобиль: {anketa.get('stamp', '')} {anketa.get('model_car', '')}\n"
            f"Цена: {anketa.get('price', 0):,} руб",
            icon='question'
        )
        
        if result:
            self.perform_accept_anketa(anketa['id'])

    def perform_accept_anketa(self, anketa_id):
        """Принять анкету (покупка автомобиля у пользователя)"""
        try:
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/admin/anketi/{anketa_id}/accept", headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Анкета принята, автомобиль добавлен в базу!")
                self.show_admin_anketi_management()
            else:
                error_msg = response.json().get("detail", "Ошибка принятия анкеты")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")


    def create_scrollable_frame(self, parent):
        """Скроллируемая область"""
        container = tk.Frame(parent, bg=self.colors['light'])
        
        canvas = tk.Canvas(container, bg=self.colors['light'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors['light'])

        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def configure_scroll_region(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(canvas_window, width=canvas.winfo_width())
        
        def on_mousewheel(event):
            current_pos = canvas.yview()

            if canvas.bbox("all"):
                content_height = canvas.bbox("all")[3]
                canvas_height = canvas.winfo_height()
                if content_height <= canvas_height:
                    return

                max_scroll = (content_height - canvas_height) / content_height

                scroll_direction = -1 if event.delta > 0 else 1
                scroll_amount = 0.1 

                new_pos = current_pos[0] + scroll_direction * scroll_amount

                if new_pos < 0:
                    new_pos = 0
                elif new_pos > max_scroll:
                    new_pos = max_scroll

                canvas.yview_moveto(new_pos)

        scrollable_frame.bind("<Configure>", configure_scroll_region)
        canvas.bind("<Configure>", configure_scroll_region)

        def bind_scroll(event):
            on_mousewheel(event)

        canvas.bind("<MouseWheel>", bind_scroll)
        scrollable_frame.bind("<MouseWheel>", bind_scroll)

        def bind_scroll_to_children(widget):
            widget.bind("<MouseWheel>", bind_scroll)
            for child in widget.winfo_children():
                bind_scroll_to_children(child)
        
        bind_scroll_to_children(scrollable_frame)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        return container, scrollable_frame, canvas, bind_scroll_to_children


    def show_stamp_management(self):
        """Управление марками автомобилей для администратора"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Управление марками автомобилей", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        management_header = tk.Label(card, text="🏷️ Управление марками автомобилей", 
                                    bg=self.colors['light'], fg=self.colors['dark'],
                                    font=('Arial', 14, 'bold'), anchor='w')
        management_header.pack(fill=tk.X, pady=(0, 15))

        actions_frame = tk.Frame(card, bg=self.colors['light'])
        actions_frame.pack(fill=tk.X, pady=10)

        buttons_container = tk.Frame(actions_frame, bg=self.colors['light'])
        buttons_container.pack(fill=tk.X)
        
        create_btn = ttk.Button(buttons_container, text="➕ Добавить марку",
                            style='Success.TButton',
                            command=self.show_create_stamp_form,
                            width=20)
        create_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        edit_btn = ttk.Button(buttons_container, text="✏️ Редактировать",
                            style='Accent.TButton',
                            command=self.show_edit_stamp_form,
                            width=20)
        edit_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        delete_btn = ttk.Button(buttons_container, text="🗑️ Удалить",
                            style='Danger.TButton',
                            command=self.confirm_delete_stamp,
                            width=20)
        delete_btn.pack(side=tk.LEFT, ipady=8)

        self.selected_stamp_info = tk.Label(actions_frame, 
                                        text="Выберите марку из таблицы",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 10), anchor='w')
        self.selected_stamp_info.pack(fill=tk.X, pady=(10, 0))

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        table_frame = tk.Frame(card, bg=self.colors['light'])
        table_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        table_header = tk.Label(table_frame, text="Список марок автомобилей:",
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 12, 'bold'), anchor='w')
        table_header.pack(fill=tk.X, pady=(0, 10))

        self.load_and_display_stamps_table(table_frame)

    def load_and_display_stamps_table(self, parent_frame):
        """Загрузить и отобразить таблицу марок с Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/stamps/", headers=headers)
            
            if response.status_code == 200:
                stamps = response.json()
                
                if not stamps:
                    no_stamps_label = tk.Label(parent_frame, 
                                            text="В системе нет марок автомобилей",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_stamps_label.pack(pady=20)
                    return

                self.create_stamps_treeview(parent_frame, stamps)
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки марок")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_stamps_treeview(self, parent, stamps):
        """Создать Treeview для отображения марок"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.stamps_tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Название марки'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            height=12
        )

        v_scrollbar.config(command=self.stamps_tree.yview)

        self.stamps_tree.heading('ID', text='ID', anchor=tk.CENTER)
        self.stamps_tree.heading('Название марки', text='Название марки', anchor=tk.W)

        self.stamps_tree.column('ID', width=80, minwidth=60)
        self.stamps_tree.column('Название марки', width=300, minwidth=200)

        self.stamps_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for stamp in sorted(stamps, key=lambda x: x['id']):
            self.stamps_tree.insert(
                '', 
                tk.END, 
                values=(
                    stamp['id'],
                    stamp['stamp']
                )
            )

        self.stamps_tree.bind('<<TreeviewSelect>>', self.on_stamp_select)

        self.configure_treeview_style()

    def on_stamp_select(self, event):
        """Обработчик выбора марки в Treeview"""
        selected_items = self.stamps_tree.selection()
        if selected_items:
            item = selected_items[0]
            stamp_data = self.stamps_tree.item(item, 'values')
            
            self.selected_stamp = {
                'id': int(stamp_data[0]),
                'stamp': stamp_data[1]
            }
            
            self.selected_stamp_info.config(
                text=f"Выбрана: {self.selected_stamp['stamp']} (ID: {self.selected_stamp['id']})",
                fg=self.colors['dark']
            )

    def show_create_stamp_form(self):
        """Показать форму создания марки"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Добавление марки автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению марками",
                            style='Secondary.TButton',
                            command=self.show_stamp_management,
                            width=35)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        field_container = tk.Frame(form_frame, bg=self.colors['light'])
        field_container.pack(fill=tk.X, pady=20)
        
        lbl = tk.Label(field_container, text="Название марки", bg=self.colors['light'], 
                    fg=self.colors['dark'], font=('Arial', 12, 'bold'),
                    anchor='w')
        lbl.pack(fill=tk.X, pady=(0, 10))
        
        self.create_stamp_entry = ttk.Entry(field_container, font=('Arial', 12))
        self.create_stamp_entry.pack(fill=tk.X, pady=2, ipady=8)
        
        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Добавить марку", 
                            style='Success.TButton', 
                            command=self.perform_create_stamp)
        save_btn.pack(fill=tk.X, pady=6, ipady=10)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_stamp_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=8)

    def show_edit_stamp_form(self):
        """Показать форму редактирования марки"""
        if not hasattr(self, 'selected_stamp') or not self.selected_stamp:
            messagebox.showwarning("Внимание", "Сначала выберите марку из таблицы")
            return
        
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование марки автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению марками",
                            style='Secondary.TButton',
                            command=self.show_stamp_management,
                            width=35)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        info_frame = tk.Frame(form_frame, bg=self.colors['light'])
        info_frame.pack(fill=tk.X, pady=(0, 20))
        
        info_text = f"Редактирование марки:\nID: {self.selected_stamp['id']}"
        info_label = tk.Label(info_frame, text=info_text,
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 11, 'bold'), anchor='w', justify=tk.LEFT)
        info_label.pack(fill=tk.X)

        field_container = tk.Frame(form_frame, bg=self.colors['light'])
        field_container.pack(fill=tk.X, pady=20)
        
        lbl = tk.Label(field_container, text="Название марки", bg=self.colors['light'], 
                    fg=self.colors['dark'], font=('Arial', 12, 'bold'),
                    anchor='w')
        lbl.pack(fill=tk.X, pady=(0, 10))
        
        self.edit_stamp_entry = ttk.Entry(field_container, font=('Arial', 12))
        self.edit_stamp_entry.insert(0, self.selected_stamp['stamp'])
        self.edit_stamp_entry.pack(fill=tk.X, pady=2, ipady=8)

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                            style='Success.TButton', 
                            command=self.perform_edit_stamp)
        save_btn.pack(fill=tk.X, pady=6, ipady=10)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_stamp_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=8)

    def perform_create_stamp(self):
        """Создать новую марку"""
        try:
            stamp_name = self.create_stamp_entry.get().strip()
            
            if not stamp_name:
                messagebox.showerror("Ошибка", "Введите название марки")
                return
            
            data = {"stamp": stamp_name}
            
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/admin/stamps/", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Марка успешно создана!")
                self.show_stamp_management()
            else:
                error_msg = response.json().get("detail", "Ошибка создания марки")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def perform_edit_stamp(self):
        """Редактировать марку"""
        try:
            stamp_name = self.edit_stamp_entry.get().strip()
            
            if not stamp_name:
                messagebox.showerror("Ошибка", "Введите название марки")
                return
            
            data = {"stamp": stamp_name}
            
            headers = {"token": self.auth_token}
            stamp_id = self.selected_stamp['id']
            response = requests.put(f"{API_BASE_URL}/admin/stamps/{stamp_id}", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Марка успешно обновлена!")
                self.show_stamp_management()
            else:
                error_msg = response.json().get("detail", "Ошибка обновления марки")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_delete_stamp(self):
        """Подтверждение удаления марки"""
        if not hasattr(self, 'selected_stamp') or not self.selected_stamp:
            messagebox.showwarning("Внимание", "Сначала выберите марку из таблицы")
            return
        
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить марку?\n\n"
            f"ID: {self.selected_stamp['id']}\n"
            f"Название: {self.selected_stamp['stamp']}",
            icon='warning'
        )
        
        if result:
            self.perform_delete_stamp()

    def perform_delete_stamp(self):
        """Удалить марку"""
        try:
            headers = {"token": self.auth_token}
            stamp_id = self.selected_stamp['id']
            response = requests.delete(f"{API_BASE_URL}/admin/stamps/{stamp_id}", 
                                    headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Марка успешно удалена!")
                if hasattr(self, 'selected_stamp'):
                    self.selected_stamp = None
                self.show_stamp_management()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления марки")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def show_model_management(self):
        """Управление моделями автомобилей для администратора"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Управление моделями автомобилей", 
                                style='Header.TLabel')
        header_label.pack(pady=(10, 5))

        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад в главное меню",
                            style='Secondary.TButton',
                            command=self.show_main_menu,
                            width=25)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        management_header = tk.Label(card, text="🚙 Управление моделями автомобилей", 
                                    bg=self.colors['light'], fg=self.colors['dark'],
                                    font=('Arial', 14, 'bold'), anchor='w')
        management_header.pack(fill=tk.X, pady=(0, 15))

        actions_frame = tk.Frame(card, bg=self.colors['light'])
        actions_frame.pack(fill=tk.X, pady=10)

        buttons_container = tk.Frame(actions_frame, bg=self.colors['light'])
        buttons_container.pack(fill=tk.X)
        
        create_btn = ttk.Button(buttons_container, text="➕ Добавить модель",
                            style='Success.TButton',
                            command=self.show_create_model_form,
                            width=20)
        create_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        edit_btn = ttk.Button(buttons_container, text="✏️ Редактировать",
                            style='Accent.TButton',
                            command=self.show_edit_model_form,
                            width=20)
        edit_btn.pack(side=tk.LEFT, padx=(0, 10), ipady=8)
        
        delete_btn = ttk.Button(buttons_container, text="🗑️ Удалить",
                            style='Danger.TButton',
                            command=self.confirm_delete_model,
                            width=20)
        delete_btn.pack(side=tk.LEFT, ipady=8)

        self.selected_model_info = tk.Label(actions_frame, 
                                        text="Выберите модель из таблицы",
                                        bg=self.colors['light'], fg='#7f8c8d',
                                        font=('Arial', 10), anchor='w')
        self.selected_model_info.pack(fill=tk.X, pady=(10, 0))

        separator = ttk.Separator(card, orient='horizontal')
        separator.pack(fill=tk.X, pady=20)

        table_frame = tk.Frame(card, bg=self.colors['light'])
        table_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        table_header = tk.Label(table_frame, text="Список моделей автомобилей:",
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 12, 'bold'), anchor='w')
        table_header.pack(fill=tk.X, pady=(0, 10))

        self.load_and_display_models_table(table_frame)

    def load_and_display_models_table(self, parent_frame):
        """Загрузить и отобразить таблицу моделей с Treeview"""
        try:
            headers = {"token": self.auth_token}
            response = requests.get(f"{API_BASE_URL}/admin/models/", headers=headers)
            
            if response.status_code == 200:
                models = response.json()
                
                if not models:
                    no_models_label = tk.Label(parent_frame, 
                                            text="В системе нет моделей автомобилей",
                                            bg=self.colors['light'], fg='#7f8c8d',
                                            font=('Arial', 11))
                    no_models_label.pack(pady=20)
                    return

                self.create_models_treeview(parent_frame, models)
                
            else:
                error_msg = response.json().get("detail", "Ошибка загрузки моделей")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def create_models_treeview(self, parent, models):
        """Создать Treeview для отображения моделей"""
        tree_frame = tk.Frame(parent, bg=self.colors['light'])
        tree_frame.pack(fill=tk.BOTH, expand=True)

        v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.models_tree = ttk.Treeview(
            tree_frame,
            columns=('ID', 'Название модели'),
            show='headings',
            yscrollcommand=v_scrollbar.set,
            height=12
        )

        v_scrollbar.config(command=self.models_tree.yview)

        self.models_tree.heading('ID', text='ID', anchor=tk.CENTER)
        self.models_tree.heading('Название модели', text='Название модели', anchor=tk.W)

        self.models_tree.column('ID', width=80, minwidth=60)
        self.models_tree.column('Название модели', width=300, minwidth=200)

        self.models_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for model in sorted(models, key=lambda x: x['id']):
            self.models_tree.insert(
                '', 
                tk.END, 
                values=(
                    model['id'],
                    model['model_car']
                )
            )

        self.models_tree.bind('<<TreeviewSelect>>', self.on_model_select)

        self.configure_treeview_style()

    def on_model_select(self, event):
        """Обработчик выбора модели в Treeview"""
        selected_items = self.models_tree.selection()
        if selected_items:
            item = selected_items[0]
            model_data = self.models_tree.item(item, 'values')
            
            self.selected_model = {
                'id': int(model_data[0]),
                'model_car': model_data[1]
            }
            
            self.selected_model_info.config(
                text=f"Выбрана: {self.selected_model['model_car']} (ID: {self.selected_model['id']})",
                fg=self.colors['dark']
            )

    def show_create_model_form(self):
        """Показать форму создания модели"""
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Добавление модели автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению моделями",
                            style='Secondary.TButton',
                            command=self.show_model_management,
                            width=35)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        field_container = tk.Frame(form_frame, bg=self.colors['light'])
        field_container.pack(fill=tk.X, pady=20)
        
        lbl = tk.Label(field_container, text="Название модели", bg=self.colors['light'], 
                    fg=self.colors['dark'], font=('Arial', 12, 'bold'),
                    anchor='w')
        lbl.pack(fill=tk.X, pady=(0, 10))
        
        self.create_model_entry = ttk.Entry(field_container, font=('Arial', 12))
        self.create_model_entry.pack(fill=tk.X, pady=2, ipady=8)

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Добавить модель", 
                            style='Success.TButton', 
                            command=self.perform_create_model)
        save_btn.pack(fill=tk.X, pady=6, ipady=10)

        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_model_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=8)

    def show_edit_model_form(self):
        """Показать форму редактирования модели"""
        if not hasattr(self, 'selected_model') or not self.selected_model:
            messagebox.showwarning("Внимание", "Сначала выберите модель из таблицы")
            return
        
        self.clear_window()
        
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True)
        
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        header_label = ttk.Label(header_frame, text="Редактирование модели автомобиля", 
                            style='Header.TLabel')
        header_label.pack(pady=(10, 5))
        
        back_frame = tk.Frame(header_frame, bg=self.colors['background'])
        back_frame.pack(fill=tk.X, pady=(0, 10))
        
        back_btn = ttk.Button(back_frame, text="← Назад к управлению моделями",
                            style='Secondary.TButton',
                            command=self.show_model_management,
                            width=35)
        back_btn.pack(ipady=8, anchor='center')

        card = self.create_card_frame(main_frame)
        card.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        form_frame = tk.Frame(card, bg=self.colors['light'])
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        info_frame = tk.Frame(form_frame, bg=self.colors['light'])
        info_frame.pack(fill=tk.X, pady=(0, 20))
        
        info_text = f"Редактирование модели:\nID: {self.selected_model['id']}"
        info_label = tk.Label(info_frame, text=info_text,
                            bg=self.colors['light'], fg=self.colors['dark'],
                            font=('Arial', 11, 'bold'), anchor='w', justify=tk.LEFT)
        info_label.pack(fill=tk.X)

        field_container = tk.Frame(form_frame, bg=self.colors['light'])
        field_container.pack(fill=tk.X, pady=20)
        
        lbl = tk.Label(field_container, text="Название модели", bg=self.colors['light'], 
                    fg=self.colors['dark'], font=('Arial', 12, 'bold'),
                    anchor='w')
        lbl.pack(fill=tk.X, pady=(0, 10))
        
        self.edit_model_entry = ttk.Entry(field_container, font=('Arial', 12))
        self.edit_model_entry.insert(0, self.selected_model['model_car'])
        self.edit_model_entry.pack(fill=tk.X, pady=2, ipady=8)

        btn_frame = tk.Frame(card, bg=self.colors['light'])
        btn_frame.pack(fill=tk.X, pady=(20, 10))

        save_btn = ttk.Button(btn_frame, text="💾 Сохранить изменения", 
                            style='Success.TButton', 
                            command=self.perform_edit_model)
        save_btn.pack(fill=tk.X, pady=6, ipady=10)
        
        cancel_btn = ttk.Button(btn_frame, text="❌ Отмена", 
                            style='Secondary.TButton', 
                            command=self.show_model_management)
        cancel_btn.pack(fill=tk.X, pady=6, ipady=8)

    def perform_create_model(self):
        """Создать новую модель"""
        try:
            model_name = self.create_model_entry.get().strip()
            
            if not model_name:
                messagebox.showerror("Ошибка", "Введите название модели")
                return
            
            data = {"model_car": model_name}
            
            headers = {"token": self.auth_token}
            response = requests.post(f"{API_BASE_URL}/admin/models/", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Модель успешно создана!")
                self.show_model_management()
            else:
                error_msg = response.json().get("detail", "Ошибка создания модели")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def perform_edit_model(self):
        """Редактировать модель"""
        try:
            model_name = self.edit_model_entry.get().strip()
            
            if not model_name:
                messagebox.showerror("Ошибка", "Введите название модели")
                return

            data = {"model_car": model_name}

            headers = {"token": self.auth_token}
            model_id = self.selected_model['id']
            response = requests.put(f"{API_BASE_URL}/admin/models/{model_id}", json=data, headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Модель успешно обновлена!")
                self.show_model_management()
            else:
                error_msg = response.json().get("detail", "Ошибка обновления модели")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

    def confirm_delete_model(self):
        """Подтверждение удаления модели"""
        if not hasattr(self, 'selected_model') or not self.selected_model:
            messagebox.showwarning("Внимание", "Сначала выберите модель из таблицы")
            return
        
        result = messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить модель?\n\n"
            f"ID: {self.selected_model['id']}\n"
            f"Название: {self.selected_model['model_car']}",
            icon='warning'
        )

        if result:
            self.perform_delete_model()

    def perform_delete_model(self):
        """Удалить модель"""
        try:
            headers = {"token": self.auth_token}
            model_id = self.selected_model['id']
            response = requests.delete(f"{API_BASE_URL}/admin/models/{model_id}", 
                                    headers=headers)
            
            if response.status_code == 200:
                messagebox.showinfo("Успех", "Модель успешно удалена!")
                if hasattr(self, 'selected_model'):
                    self.selected_model = None
                self.show_model_management()
            else:
                error_msg = response.json().get("detail", "Ошибка удаления модели")
                messagebox.showerror("Ошибка", error_msg)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Ошибка", f"Ошибка подключения: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = CarTradingApp(root)
    root.mainloop()