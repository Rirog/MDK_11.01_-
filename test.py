from docx import Document
from docx.shared import Pt, RGBColor
import os
file_path = 'draft.docx'
if os.path.exists(file_path):
    os.remove(file_path)

doc = Document()

title = doc.add_heading(level=1)
t  = title.add_run("Текст черный")
t.font.color.rgb = RGBColor(0, 0, 0)
text  = doc.add_paragraph()
te = text.add_run("Это основной текст документа с важной информацией.")
te.font.name = 'Times New Roman'
te.font.size = Pt(12)

doc.add_heading("Подзаголовок", level=2)
doc.add_paragraph("допольнительный текст подзаголовком.")

doc.save('draft.docx')

print("Содержимое документа:")
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        print(f"- {paragraph.text}")

print("\nЗаголовки в документе:")
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading'):
        print(f"Заголовок уровня {paragraph.style.name[-1]}: {paragraph.text}")